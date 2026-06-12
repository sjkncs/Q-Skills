"""Generate 2-page cheat sheet v7 - complete with all PPT content + gap analysis additions."""
import docx, re, os

def extract_all(path):
    doc = docx.Document(path)
    paras = [(p.style.name if p.style else 'N', p.text.strip()) for p in doc.paragraphs]
    tables = [[[c.text.strip() for c in r.cells] for r in tb.rows] for tb in doc.tables]
    return paras, tables

def find_headings(paras, level=None):
    return [(i, s, t) for i, (s, t) in enumerate(paras) if 'Heading' in s and (level is None or level in s)]

def get_section(paras, start, end):
    return [(s, t) for s, t in paras[start+1:end] if t]

# ===== Parsers =====
def parse_mcq_answers_full(items):
    ans = []; cur_n = None; cur_letter = ''; cur_body = ''
    for s, t in items:
        m = re.match(r'(\d+)[\.\s]*答案[：:]\s*([A-D])', t)
        if m:
            if cur_n is not None: ans.append((cur_n, cur_letter, cur_body.strip()))
            cur_n = int(m.group(1)); cur_letter = m.group(2); cur_body = ''
        elif cur_n is not None: cur_body += t + ' '
    if cur_n is not None: ans.append((cur_n, cur_letter, cur_body.strip()))
    return ans

def parse_fill_answers_full(items):
    ans = []; cur_n = None; cur_ans = ''; cur_body = ''
    for s, t in items:
        m = re.match(r'(\d+)[\.\s]*答案[：:]\s*(.*)', t)
        if m:
            if cur_n is not None: ans.append((cur_n, cur_ans, cur_body.strip()))
            cur_n = int(m.group(1)); cur_ans = m.group(2).strip(); cur_body = ''
        elif cur_n is not None: cur_body += t + ' '
    if cur_n is not None: ans.append((cur_n, cur_ans, cur_body.strip()))
    return ans

def parse_fill_questions(items):
    return [(int(m.group(1)), m.group(2)) for s, t in items if (m := re.match(r'(\d+)[.\、]\s*(.*)', t))]

def parse_defn_answers(items):
    ans = []; buf = ''
    for s, t in items:
        m = re.match(r'(\d+)[.\、]\s*(.*)', t)
        if m:
            if buf: ans.append(buf)
            buf = m.group(2)
        else: buf += ' ' + t
    if buf: ans.append(buf)
    return ans

def parse_essay_full(items):
    ans = []; cur_n = None; cur_title = ''; cur_body = ''
    for s, t in items:
        m = re.match(r'(\d+)[.\、]\s*(.*)', t)
        if m and not t.startswith('核心') and not t.startswith('关键'):
            if cur_n is not None: ans.append((cur_n, cur_title, cur_body.strip()))
            cur_n = int(m.group(1)); cur_title = m.group(2); cur_body = ''
        elif cur_n is not None: cur_body += t + ' '
    if cur_n is not None: ans.append((cur_n, cur_title, cur_body.strip()))
    return ans

def parse_numbered_qs(items):
    return [(int(m.group(1)), m.group(2)) for s, t in items if (m := re.match(r'(\d+)[.\、]\s*(.*)', t))]

# ===== Build knowledge map HTML (Page 2) =====
def build_knowledge_map():
    chapters = [
        ("第1章 绪论", [
            ("1.1 具身智能含义与发展", [
                "具身智能/Embodied Intelligence: 智能体通过身体与环境交互实现智能行为",
                "具身感知/Perception: 多模态传感器获取环境信息(视觉/触觉/力觉)",
                "具身推理/Reasoning: 基于物理世界约束的因果推理与规划",
                "具身执行/Execution: 运动控制、抓取操作、导航避障",
                "空间智能vs具身智能: 空间=数字空间软件(API/文本),部署2-4周; 具身=物理载体(机器人/无人机),研发6-12+月",
                "智能机器人案例: 波士顿动力Atlas/特斯拉Optimus/Figure 01",
                "具身智能应用: 工业制造、家庭服务、医疗辅助、仓储物流",
                "未来方向: 通用具身智能、人机协作、群体智能"
            ]),
            ("1.2 大模型基本概念", [
                "大规模语言模型LLM: 基于Transformer的大规模参数神经网络",
                "概率分布建模: P(w1,w2,...,wm) 对词序列联合概率建模",
                "|V|^m爆炸: 7万词条×20词长→7.98×10^96种可能(天文数字)",
                "N元语言模型/N-gram: 马尔可夫假设,仅依赖前n-1个词",
                "神经语言模型: 用神经网络替代计数,连续空间表示",
                "独热编码One-Hot: 稀疏高维表示,维度=词表大小",
                "循环神经网络RNN: 输入层→隐藏层→输出层,时序建模",
                "自监督学习: 无需标注,从文本本身构造学习信号(下一词预测)",
                "预训练语言模型: 大规模无标注数据训练通用表示"
            ]),
            ("1.3 大语言模型发展历程(三阶段)", [
                "基础模型阶段(2018-2021): Transformer(2017)→BERT/GPT-1(2018)→GPT-2(1.5B)/T5(11B,2019)→GPT-3(175B,2020)",
                "能力探索阶段(2019-2022): GPT-2零样本→GPT-3 ICL少样本→InstructGPT(SFT+RL,2022)",
                "突破阶段(2022.11 ChatGPT起): GPT-4(2023.3,SAT数学95%/高考英语92.5%),Bard/文心/星火/ChatGLM/MOSS",
                "GPT-4: 2023.3.14发布,支持视觉输入,多模态理解,基准考试超88%",
                "MaaS(Model as a Service): GPT-3后出现的模型即服务范式,提示学习+指令微调"
            ]),
            ("1.4 大语言模型构建流程(四阶段)", [
                "预训练Pre-training: 数千亿词+1000+GPU+月级(代表:GPT-3/LLaMA/PaLM)",
                "有监督微调SFT: 数万指令+1-100GPU+天级(代表:MOSS/ChatGLM/Vicuna)",
                "奖励建模RM: 百万对比对+1-100GPU+天级,Bradley-Terry偏好建模",
                "强化学习RLHF: 十万指令+1-100GPU+天级,PPO策略优化",
            ]),
        ]),
        ("第2章 大语言模型基础", [
            ("2.1 Transformer结构", [
                "嵌入表示层Embedding: 词→稠密向量,维度d_model,可学习参数",
                "位置编码PE: sin/cos正余弦函数,PE(pos,2i)=sin(pos/10000^(2i/d))",
                "自注意力Self-Attention: QKV点积+缩放因子√dk+softmax,全局依赖",
                "多头注意力Multi-Head: h组并行注意力,不同子空间,Wi^Q/Wi^K/Wi^V/W^O",
                "前馈层FFN: 两层全连接+ReLU,d_ff通常=4×d_model",
                "残差连接: x_{l+1}=x_l+f(x_l),缓解梯度消失",
                "层归一化LayerNorm: 稳定激活值分布,应用于残差后",
                "编码器Encoder: 双向自注意力+FFN×N层,理解任务",
                "解码器Decoder: 掩码自注意力+交叉注意力+FFN×N层,生成任务",
                "交叉注意力: Q来自解码器,K/V来自编码器输出"
            ]),
            ("2.2 生成式预训练GPT", [
                "GPT模型: 多层Transformer解码器(去掉交叉注意力),单向语言模型",
                "无监督预训练: 最大化L=ΣlogP(wt|w<t),自回归左到右生成",
                "有监督微调: 取最后一层最后一词h^(L)_n,经FC+softmax预测标签",
                "灾难性遗忘对策: 混合损失L=L_sft+λ·L_pretrain",
                "语料预处理: 文档级拼接+随机打乱+去重+质量过滤",
                "句子补全: 给定前缀,自回归逐token生成直至EOS"
            ]),
            ("2.3 大语言模型的结构(LLaMA系列)", [
                "LLaMA架构: 解码器only, RMSNorm+SwiGLU+RoPE+GQA",
                "RMSNorm: 替代LayerNorm,省去均值计算,a/RMS(a)·γ",
                "SwiGLU激活: (xW1⊙σ(xW3))W2,优于ReLU/GELU,LLaMA标配",
                "Swish: x·σ(βx),β→0线性,β→∞趋近ReLU,HF用SiLU替代",
                "RoPE旋转位置编码: q·e^(imθ),绝对编码实现相对位置,θ=10000^(-2i/d)",
                "稀疏注意力: 五种类型(Global/Band/Dilated/Random/Block Local),降低O(n²)复杂度",
                "Reformer LSH: h(x)=argmax([xR;-xR]),R∈R^(dk×b/2),局部敏感哈希近似注意力",
                "Routing Transformer: K-means聚类Q和K,中心向量滑动平均更新,同簇内交互",
                "FlashAttention: 分块计算,S/P/O不写入全局内存,SRAM 228KB/块",
                "多查询注意力MQA: 共享K/V仅保留Q,5%数据微调即有效,减少KV缓存(Falcon/StarCoder采用)",
                "分组查询GQA: 介于MHA和MQA之间,每组共享K/V,LLaMA-2 70B采用",
                "GPT-4 MoE架构: 1.8T总参,16专家×1110亿,路由2个,共享注意力550亿"
            ]),
        ]),
        ("第3章 预训练数据", [
            ("3.1 数据来源", [
                "通用数据: CommonCrawl(45TB原始→570GB过滤→500B词元),Wikipedia(多语言),Google索引130万亿网页",
                "专业领域数据: 学术论文(arXiv/PubMed),代码(GitHub),法律/医学/金融",
                "多语言规模: BLOOM训练46种语言,PaLM训练122种语言",
                "科学数据表示: 公式用LaTeX,化学结构用SMILES,蛋白质用单/三字母编码",
                "数据多样性: 多语言、多领域、多风格,避免单一数据源偏差"
            ]),
            ("3.2 数据处理流水线", [
                "质量过滤: 分类器打分+规则过滤(长度/特殊字符/困惑度),保留高质量文本",
                "冗余去除: MinHash/SimHash近似去重,布隆过滤器O(k)时间复杂度",
                "后缀数组去重: RefinedWeb k≥50过滤公共子串; Wiki-40B(4GB)→140秒; C4(350GB)→12小时",
                "隐私消除PII: 正则匹配+NER识别姓名/邮箱/电话/地址,替换为占位符",
                "词元切分Tokenization: 文本→子词序列,BPE/WordPiece/SentencePiece/Unigram",
                "BPE: 全词→单字符+</w>→统计频率→合并最高频→重复(LLaMA词表32K)",
                "WordPiece: Score=Cnt(pair)/(Cnt(a)·Cnt(b)),最大化似然",
                "SentencePiece: 直接对原始文本操作,不依赖预分词,支持BPE/Unigram",
                "Unigram(T5/mBART): 从大候选集迭代删除token,EM算法+Viterbi动态规划,与BPE反向操作"
            ]),
            ("3.3 数据影响分析", [
                "数据规模→Scaling Law: 三因素幂律(参数N×数据D×计算C),效果随指数线性提高",
                "Chinchilla验证: 70B+1.4T词元等比缩放,优于Gopher(280B)/GPT-3(175B)",
                "数据质量-重复: 0.1%数据×100次重复→800M模型降至400M水平;HP复制11次→3%严重影响",
                "PaLM记忆率: 见1次→0.75%记忆率; 见500+次→>40%记忆率",
                "数据时间性: Google训练28个1.5B模型(C4 2013/2016/2019/2022),训练-测试时间不匹配影响性能",
                "数据质量-噪声: 低质量数据污染训练,需过滤+清洗",
                "数据多样性: 多领域混合提高泛化能力; Gopher消融:最优=10%C4+50%MassiveWeb+30%Books+10%News",
                "LLaMA-2: 2T词元训练; Baichuan-2: 2.6T词元; 训练文件~10TB"
            ]),
            ("3.4 开源数据集", [
                "Pile: 825GB,22个数据源混合(EleutherAI)",
                "ROOTS: 1.6T token,46个子数据集(Google)",
                "RefinedWeb: 5T token,高质量CC子集(Falcon使用)",
                "SlimPajama: 627B token,RedPajama清洗版(CerebrasAI)"
            ]),
        ]),
        ("第4章 分布式训练", [
            ("4.1 分布式训练三墙", [
                "计算墙: H100 FP16=2000TFLOPs vs GPT-3=314ZFLOPs,差8个数量级",
                "显存墙: GPT-3 FP32=700GB,FP16=350GB vs H100仅80GB",
                "通信墙: 128副本×700GB梯度=89.6TB/迭代,IB链路≤800Gbps",
                "模型增速: AlexNet(2013)→PaLM(2022)每18个月增长56倍",
                "总训练速度=f(单设备速度×设备总量×多设备加速比)"
            ]),
            ("4.2 并行策略详解", [
                "数据并行DP: N样本/M设备,每设备N/M,G_avg=ΣGi/M,AllReduce同步梯度",
                "流水线并行PP: 模型按层切分到不同设备,微批次流水线减少气泡",
                "GPipe: 将mini-batch拆为micro-batch,F1→F11/F12/F13/F14,减少流水线气泡",
                "Megatron-LM 1F1B: 1次前向+1次反向交替,三阶段(warmup/forward-backward/cooldown)",
                "1F1B交错调度: micro-batch数=流水线阶段整数倍,每设备处理多个模型块",
                "张量并行TP: 矩阵Y=XA按行切(A=[A1;A2])或按列切(A=[A1|A2])",
                "FFN张量并行策略: 第1个FC按列切块,第2个FC按行切块,省去中间通信",
                "Embedding TP: 64000×5120×4B≈1.25GB×2(含梯度)=2.5GB",
                "混合并行/3D并行: DP+PP+TP组合,适应超大规模模型",
                "ZeRO Stage1: Adam状态分区→内存1/4; Stage2: +梯度分区→1/8; Stage3: +参数→线性",
                "ZeRO-Infinity: Stage3+CPU/NVMe卸载,可训练万亿参数模型",
                "梯度累积: eff_batch=micro_batch×accum_steps,等效放大batch不增显存",
                "激活重计算: 前向时丢弃中间激活,反向时重算,用计算换内存",
                "混合精度: 参数2Φ+梯度2Φ+Adam(4+4+4)Φ=16Φ字节(Φ=参数量)"
            ]),
            ("4.3 集群架构与网络", [
                "HPC集群: GPU节点+高速互联(InfiniBand/NVLink),计算/存储/管理网络分离",
                "参数服务器: 中心节点存储参数,worker节点计算梯度,瓶颈在PS带宽",
                "去中心化: Ring AllReduce,每步通信量2(N-1)/N·M,带宽最优",
                "网络带宽: PCIe5.0=128GB/s, IB=200-400Gbps, DGX=1.6Tb, HGX=3.2Tb",
                "DGX GH200(2023.5): 256个GH200芯片,NVLink互联,144TB共享内存",
                "<30B模型无需张量并行,现有DP+PP框架即可多节点多GPU训练",
                "分布式Softmax: softmax(x)_p=exp(x_p-max)/Σexp(x_q-max),仅需三次通信"
            ]),
            ("4.4 DeepSpeed与训练实践", [
                "ZeRO-3配置: stage3参数,分片策略,通信重叠",
                "CPU Offload: 将优化器状态卸载到CPU内存,缓解GPU显存压力",
                "NVMe Offload: 进一步卸载到NVMe SSD,容量更大但速度更慢",
                "LLaMA训练资源: 7B=82K, 13B=135K, 65B=1022K GPUh; 70B=172万GPUh",
                "1-bit Adam: 通信量仅原始Adam的1/5,达到类似收敛率",
                "Adam超参: β1=0.9,β2=0.95,ε=1e-8,clip=1.0,wd=0.1",
                "LR调度: 热身0.1-0.5%→峰值5e-5→余弦衰减至10%; 批次: LLaMA-2=4M词元"
            ]),
        ]),
        ("第5章 有监督微调", [
            ("5.1 提示学习与语境学习", [
                "提示学习Prompt Learning: 设计模板引导模型,无需更新参数",
                "语境学习ICL: 在prompt中提供示例,模型推断任务模式",
                "Few-Shot: 提供K个示例(K通常1-64); Zero-Shot: 仅给任务描述",
                "提示模板设计: 任务描述+输入格式+输出要求,影响模型表现",
                "ICL关键: 示例选择(多样性/相关性)+示例顺序+格式一致性",
            ]),
            ("5.2 高效模型微调PEFT", [
                "LoRA: W'=W0+BA, B(d×r)随机初始化, A(r×d)零初始化, scale=α/r",
                "LoRA效果: r=4仅Q/V矩阵, 350GB→35MB(万分之一), GPU 1.2TB→350GB, 速度↑25%",
                "AdaLoRA: ΔW=PΓQ, SVD参数化, 自适应分配秩, 重要性分数=奇异值",
                "QLoRA: NF4量化(基于N(0,1)分位数,4bit=16分位), 65B模型仅需48GB显存",
                "推理合并: W=W0+BA, 无额外推理延迟"
            ]),
            ("5.3 上下文窗口扩展", [
                "位置编码外推: ALiBi线性偏置,不在Embedding加PE,softmax后加-m|i-j|",
                "ALiBi斜率序列: m=2^(-8/n),2^(-16/n),...,几何序列,头数越多衰减越快",
                "位置编码插值PI: f(x,m·L_pretrain/L_extend),将位置缩放到预训练范围内",
                "NTK-Aware Scaling: 非均匀缩放不同频率,低频外推高频插值",
                "RoPE扩展效果: LLaMA窗口扩至32768,仅需在Pile上微调1000步"
            ]),
            ("5.4 指令数据构建", [
                "手动构建: 专家撰写高质量指令,LIMA数据集(仅1000条≈数十倍数据效果,Less is More)",
                "自动构建Self-Instruct: 175条种子指令→GPT生成→过滤→扩充,迭代构建大规模数据",
                "ROUGE-L过滤: 新指令与池中任何指令的ROUGE-L<0.7才可加入,保证多样性",
                "LLaMA默认max_input=2048 token; T5 Bias PE可从512外推至~600",
                "开源指令集: Alpaca(52K)/ShareGPT(70K)/OpenAssistant/FLAN"
            ]),
            ("5.5 DeepSpeed-Chat SFT实践", [
                "代码结构: data→model→train三模块,支持DeepSpeed ZeRO",
                "数据预处理: 对话格式化→tokenize→padding→dataloader",
                "模型训练: 混合精度+梯度累积+学习率调度+检查点保存",
                "DeepSpeed配置: ZeRO Stage2/3,通信重叠,内存优化",
            ]),
        ]),
        ("第6章 强化学习", [
            ("6.1 RLHF概述", [
                "强化学习RL: 智能体通过与环境交互最大化累计奖励",
                "RL vs SFT: SFT=逐token交叉熵(per-token反馈),RL=整段奖励(whole-text反馈),RL可捕捉否定词等细微变化",
                "RLHF三阶段: SFT→奖励模型RM→PPO策略优化",
                "3H原则(InstructGPT): Helpfulness(有帮助)/Honesty(诚实)/Harmlessness(无害)",
                "三类用户输入: (a)Text-Grounded文本问答 (b)Knowledge-Seeking知识查询(RL减少幻觉) (c)Creative创意生成",
                "RL vs RLHF: 目标(环境奖励vs人类偏好)/奖励来源(环境vs人工标注)/场景(游戏/机器人vs对话/对齐)"
            ]),
            ("6.2 奖励模型", [
                "偏好数据: 对同一prompt生成多个回复,人类标注排序(chosen vs rejected),标注一致性仅60-70%",
                "Bradley-Terry模型: P(y_w≻y_l)=σ(r(x,y_w)-r(x,y_l)),基于成对比较",
                "RM训练损失: L=-logσ(r_w-r_l), 带模仿学习: L_rm=L_p+β_rm·L_lm",
                "L_lm项: 自回归语言模型损失,顶层线性层维度=|V|,防止RM偏离语言模型",
                "开源RM数据: HH-RLHF(161K)/Stanford-SHP(385K,18领域)/Summarize(179K+15K)/WebGPT(19K)"
            ]),
            ("6.3 近端策略优化PPO", [
                "策略梯度(基本): ∇J=E[R·∇log p]; 加基线: (R-b)·∇log p, b=E[R]",
                "精细奖励(Reward-to-go): 仅累计当前动作之后的奖励Σ_{t'≥t}r_{t'}",
                "折扣因子γ: Σγ^(t'-t)·r_{t'}, γ∈[0,1],越远奖励权重越低",
                "GAE优势估计: A=Q-V, L_V=E[(V_φ-Q)²], 用神经网络拟合V",
                "重要性采样: E_p[f]=E_q[f·p/q], 允许用旧策略数据更新新策略",
                "PPO-Clip: min(r·A, clip(r,1-ε,1+ε)·A), 裁剪比率防止策略更新过大",
                "PPO-Penalty: E[r·A]-β·KL(θ‖θ'), KL约束+拉格朗日自适应β",
                "Clip函数: r>1+ε→1+ε; 1-ε≤r≤1+ε→原值; r<1-ε→1-ε"
            ]),
            ("6.4 MOSS-RLHF & 红队测试", [
                "MOSS-RLHF: 基于LLaMA训练奖励模型,PPO微调对齐",
                "PPO-Max: 7个关键稳定性因素(KL惩罚/奖励裁剪/Critic损失裁剪等),确保RLHF稳定训练",
                "Reward Hacking: 模型快速提高奖励但输出无意义/重复→解法:增大KL惩罚,不能仅依赖奖励分数",
                "红队A2C损失: α·L_A2C+(1-α)·KL(p_r‖p_init), 防止RL塌陷到单一高奖励",
                "红队测试4种生成: (1)零样本 (2)随机少样本 (3)监督学习微调 (4)A2C强化学习",
                "红队测试: 对抗性提示攻击,检测模型安全边界与有害输出"
            ]),
        ]),
        ("第7章 大语言模型应用", [
            ("7.1 推理规划", [
                "思维链CoT: Let's think step by step,逐步推理提高复杂任务准确率",
                "由少至多提示: 先分解子问题,再逐个求解,最后组合",
                "自洽性Self-Consistency: 多次采样CoT,多数投票选最一致答案",
                "Zero-Shot CoT: 仅加'Let's think step by step'即可激活推理能力",
                "Auto-CoT: 人工CoT不同作者准确率差异达28.2%; Auto-CoT两阶段:K-means聚类+每簇采样",
                "Complex-CoT: 选最复杂推理样本为示范; Self-Polish: 优化歧义/低质量问题",
                "任务分解: 将复杂任务拆为可独立求解的子任务,DAG执行"
            ]),
            ("7.2 LangChain综合框架", [
                "Model I/O: Prompt模板+LLM调用+输出解析器(PydanticOutputParser)",
                "数据连接: 文档加载器(PDF/HTML/CSV)+文本分割器+向量嵌入(OpenAI/HF)",
                "向量存储: FAISS/Chroma/Pinecone,支持相似度检索,用于RAG",
                "链Chain: 顺序链SequentialChain,将多个链组合为流水线",
                "记忆Memory: 短期(对话缓冲)+长期(向量库检索),支持上下文管理",
                "智能体Agent: ReAct(思考+行动循环),工具调用(Search/Calculator/API)",
                "回调Callback: 监控链执行过程,日志记录/流式输出/token计数",
                "RAG检索增强生成: 用户问题→向量检索→注入上下文→LLM生成答案",
                "知识库问答: 文档切分→嵌入→存入向量库→检索→生成,企业级应用"
            ]),
            ("7.3 智能体Agent", [
                "LLM Agent组成: 思考模块+记忆模块+工具调用模块",
                "思考模块: ReAct(Reasoning+Acting),Thought→Action→Observation循环",
                "记忆模块: 短期(对话历史缓冲)+长期(外部存储+检索)",
                "工具调用: 函数定义→LLM选择工具→执行→结果回传",
                "多智能体辩论: 多LLM扮演不同角色,通过辩论达成共识",
                "角色扮演: 设定人物背景+性格+目标,模拟对话与协作",
                "CAMEL框架: 双Agent(用户+助手),可选任务规范Agent+评价Agent引导交互方向",
                "Agent工具: Search/Calculator/API/Code Interpreter,扩展LLM能力边界",
                "AutoGPT: 自主任务分解+规划+执行+反思,无需人类干预的Agent范式",
            ]),
            ("7.4 多模态大模型", [
                "架构: 视觉编码器ViT+线性投影+LLM(Vicuna/LLaMA),视觉→语言空间对齐",
                "GPT-4(2023.3.14): 支持视觉输入,可描述图像/解释视觉现象/从手绘草图生成HTML网站",
                "MiniGPT-4: EVA-CLIP ViT-G/14+Q-Former(BERT交叉注意力)+Vicuna-13B",
                "MiniGPT-4训练: 预训练20K步/256批/5M图文对/4×A100/10h; 微调400步/批12/1GPU/7min",
                "MiniGPT-4高质量数据: 5000图→ChatGPT过滤→保留3500; 冻结ViT+LLM仅训练投影层",
                "LLaVA: CLIP视觉编码+LLaMA,指令跟随数据微调",
                "训练策略: 预训练(大规模图文对)+微调(高质量对话数据)"
            ]),
            ("7.5 推理优化", [
                "FastServe(北大): Skip-join MLFQ调度器,k级优先级,Job Profiler定初始优先级,迭代级抢占",
                "Orca: 每批仅执行1次迭代(1 token/job),完成即离开/新任务即加入,FCFS策略,有队头阻塞",
                "KV Cache两阶段: 初始化(处理输入prompt,生成KV)→解码(仅计算新token的QKV,增量更新KV)",
                "vLLM: PagedAttention,非连续KV缓存,避免60-80%内存浪费,吞吐↑24×",
                "PagedAttention: KV缓存分页存储,类似操作系统虚拟内存,按需分配",
                "KV Cache优化: 复用已计算的K/V,避免重复计算; FastServe将低优先级KV卸载到Host内存",
                "批处理策略: 静态批/动态批/连续批处理,提高GPU利用率",
                "模型量化: FP16→INT8→INT4,减少显存+加速推理,精度-速度权衡"
            ]),
        ]),
        ("第8章 大语言模型评估", [
            ("8.1 评估概述", [
                "截至2023.8,国内外已发布120+开源/闭源大语言模型",
                "评估必要性: 模型能力涌现,传统指标不足以衡量综合能力",
                "挑战: 基准污染/评估偏差/多语言覆盖/安全对齐",
                "AI2 GENIE: 800条机器翻译人工评估成本~$80",
                "评估维度: 语言能力/推理能力/知识覆盖/安全性/多语言/对齐程度",
            ]),
            ("8.2 评估体系", [
                "任务核心评估: HELM框架(42场景/59指标/16核心场景),三维分类(任务/领域/语言)",
                "人类核心评估: AGIEval(中国高考/美国SAT/LSAT/数学竞赛/司法/公务员,GPT-4 SAT数学95%)",
                "伦理安全评估: 8类安全场景,6种指令攻击,6000+评估数据,~2800攻击样本",
                "社会偏见: CrowS-Pairs(1508例,9类偏见:种族/性别/性取向/宗教/年龄/国籍/残疾/外貌/社经)",
                "Winogender: 120人工构建句对,性别名词替换(he/she)检测性别偏见",
                "LLaMA-2安全风险三类: (1)违法犯罪 (2)有害冒犯 (3)不合格建议(医疗/金融/法律)",
                "6种指令攻击: 目标劫持/提示泄露/角色扮演/不安全主题/注入不可见内容/反向暴露",
                "红队测试Red Teaming: 4种生成方法(零样本/少样本/监督学习/A2C强化学习)",
                "垂直领域评估: 复杂推理/环境交互/特定领域(法律CUAD 500+合同/医学MultiMedQA 6数据集)"
            ]),
            ("8.3 评估方法", [
                "分类: P=TP/(TP+FP), R=TP/(TP+FN), F1=2PR/(P+R), PR曲线, BEP=P=R平衡点",
                "回归: MAE/MSE/RMSE/MAPE/MSLE/MedAE",
                "困惑度: PPL=2^H=P(s)^(-1/W), 英文n-gram: 50-1000, 交叉熵: 6-10 bit",
                "BLEU(精确率导向): BP·exp(ΣWn·logPn), BP惩罚短句",
                "ROUGE(召回率导向): ROUGE-N(n-gram匹配)/ROUGE-L(LCS最长公共子序列)",
                "三层文本质量评估: 语言层(词汇/语法/篇章)→语义层(准确性/连贯性/风格)→知识层(准确性/丰富性/一致性)",
                "McNemar: χ²=(|B-C|-1)²/(B+C), B+C<25用二项式精确检验",
                "Cohen's Kappa(2人)/Fleiss' Kappa(≥3人): Pa/Pe多评估者一致性",
                "人工评估MOS: 均分,一致性百分比; LLM-as-a-Judge: 用GPT-4评估其他模型",
                "Elo评分: 成对比较,类似国际象棋等级分,Chatbot Arena使用",
                "GPT-3数据泄漏检测: 去除与预训练集有13-gram重叠的样本,比较clean vs full; 25%基准>50%污染但影响小",
                "pass@k: 生成k个代码样本,至少1个通过测试的概率,代码评估核心指标",
            ]),
            ("8.4 评估实践", [
                "MMLU: 15858题/57学科/4选1, 划分(285dev/1531val/14042test), 众包34.5% vs 专家89.8%",
                "C-EVAL: 52学科/4难度(初高中/大学/职业), 大学覆盖13本科大类25学科, C-EVAL HARD需高级推理",
                "Chatbot Arena: 盲测对比+Elo排名, GPT-4 vs GPT-3.5胜率79%, vs LLaMA-13B胜率94%, 33K对话",
                "LLMEVAL-1: 17大类/453题, 5评分维度(正确/流畅/信息量/逻辑/无害), 5种评估方法对比",
                "数学推理: GSM8K(8500小学)/MATH(12500竞赛)/LISA(183K定理+216万证明步)/miniF2F(488 IMO级)",
                "法律: CUAD(500+合同/41条款/13000+标注), CAIL2018(260万刑事/183法条), LeCaRD(107查询+43000候选)",
                "代码评估: HumanEval(164题)/MBPP(974题)/CodeContest(竞赛级),pass@k指标",
                "多模态评估: VQAv2/GQA/TextVQA(视觉问答), MMMU(多学科多模态理解)",
            ]),
        ]),
    ]
    
    html_parts = []
    for ch_title, sections in chapters:
        html_parts.append(f'<div class="ch"><b class="cht">{ch_title}</b>')
        for sec_title, points in sections:
            if not points:
                continue
            pts = ' · '.join(points)
            html_parts.append(f'<div class="sc"><b class="shd">{sec_title}</b> {pts}</div>')
        html_parts.append('</div>')
    return '\n'.join(html_parts)

# ===== Fill-in-blank knowledge review (Page 2 lower half) =====
def build_fillblank_review():
    """PPT多样化知识点速记 - 按章节组织，7种类型混合，答案红字标注。"""
    R = lambda ans: f'<span class="r">{ans}</span>'
    chapters = [
        ("第1章 绪论", [
            f'[细节] 空间智能形成"{R("感知-决策-行动")}"闭环，依赖大语言模型、推理引擎、规划算法',
            f'[细节] 具身智能在"{R("感知→决策→行动→再感知")}"的实时闭环中持续进化',
            f'[数值] GPT-3训练中Wikipedia平均训练{R("3.4")}次epoch,CommonCrawl仅{R("0.44")}次',
            f'[数值] OPT使用{R("992")}块NVIDIA A100 80G GPU，训练近{R("2")}个月',
            f'[数值] BLOOM训练花费{R("3.5")}个月，{R("48")}节点×8块A100=共{R("384")}GPU',
            f'[细节] BLOOM节点内用{R("4个NVLink")}通信，节点间用{R("Omni-Path 100Gbps")}构建8维超立方体拓扑',
            f'[细节] BLOOM框架: {R("Megatron-LM")}提供张量并行+{R("DeepSpeed")}提供ZeRO优化器+模型流水线',
            f'[案例] Alpaca/Vicuna/MOSS/ChatGLM-6B等SFT模型达到ChatGPT {R("90%")}的效果',
            f'[对比] SFT用{R("交叉熵")}逐token调整;RL给出{R("整段")}奖励,可捕捉否定词等细微变化',
            f'[对比] RL会降低基础模型的{R("熵")},减少{R("输出多样性")};RLHF超参数众多,收敛难度大',
            f'[里程碑] {R("2019")}年大模型爆发增长;{R("2022.11")}ChatGPT发布引起全球关注',
            f'[算法] 语言模型通过{R("链式法则")}分解为逐词条件概率的乘积，降低建模复杂度',
            f'[算法] 联合概率链式分解: P(w₁,...,wm)={R("∏ᵢP(wᵢ|w₁,...,wᵢ₋₁)")},是所有自回归模型的{R("数学起点")}',
            f'[数值] 现代汉语词典第七版{R("7万")}词条×20词长→{R("7.98×10^96")}种可能',
            f'[细节] RM模型不能{R("单独")}提供给用户，仅用于对SFT模型输出进行{R("排序评估")}',
            f'[细节] SFT模型具备{R("初步指令理解")}能力,能完成开放域问答/{R("阅读理解")}/翻译/生成代码',
            f'[细节] 用户可使用{R("自然语言")}与LLM交互,实现问答/分类/{R("摘要")}/翻译/聊天等任务',
        ]),
        ("第2章 大语言模型基础", [
            f'[细节] 注意力三组件: Query({R("当前关注目标")})/Key({R("匹配元素")})/Value({R("实际内容")})',
            f'[细节] LLaMA与标准Transformer有{R("四")}处关键不同: RMSNorm+SwiGLU+RoPE+GQA',
            f'[细节] RMSNorm省去均值计算,可引入可学习{R("缩放因子gi")}和{R("偏移参数bi")}',
            f'[数值] Swish: β→0趋近{R("线性")}y=x; β→∞趋近{R("ReLU")}; β=1时光滑且{R("非单调")}',
            f'[算法] RoPE借助{R("复数")}思想,通过{R("绝对")}位置编码实现{R("相对")}位置编码',
            f'[细节] RoPE矩阵R具有{R("稀疏性")}，可用逐位{R("相乘⊗")}操作提高计算速度',
            f'[对比] 稀疏注意力分两类: 基于{R("位置")}(Global/Band等)和基于{R("内容")}(Routing/Reformer)',
            f'[算法] Routing Transformer: {R("K-means聚类")}Q和K,中心向量用{R("滑动平均")}更新',
            f'[数值] GPU显存6种: 全局/本地/{R("共享SRAM")}/寄存器/常量/纹理内存',
            f'[数值] H100全局内存{R("80GB")},访问速度{R("3.35TB/s")}; SRAM仅约{R("228KB")}/块',
            f'[细节] FlashAttention: Q/K/V分块加载到{R("SRAM")}计算,中间结果S/P/O{R("不写回")}全局内存',
            f'[细节] LLaMA采用{R("前置层归一化")}Pre-normalization: LayerNorm移到注意力层和FFN{R("之前")}',
            f'[细节] GPT-2相比GPT引入前置层归一化,残差连接调整到注意力层与FFN层{R("之后")}',
            f'[对比] Cross-Attention: Q来自{R("解码器")}当前层输出,K和V来自{R("编码器")}最终输出; 多模态对齐也用此思想',
            f'[细节] Masked Attention: 解码器生成时必须{R("避免看到未来")}词元,使用{R("掩码")}使每个位置只关注自身及之前',
            f'[数值] 标准自注意力时间和空间复杂度均为{R("O(n²d)")},随序列长度{R("平方")}增长,是长上下文的核心瓶颈',
            f'[里程碑] GPT-3后OpenAI{R("不再开源")}模型代码,ChatGPT和GPT-4架构{R("未公开")}',
            f'[案例] OPT是仿照{R("GPT-3")}架构开源的复现模型;Meta AI开源了{R("LLaMA")}',
            f'[对比] GQA({R("分组查询注意力")}): 介于MHA和MQA之间,每组{R("共享K/V")};LLaMA-2 70B采用',
            f'[数值] GPT-4推测MoE: {R("1.8T")}总参,{R("16")}专家×1110亿,路由选{R("2")}个,共享注意力550亿',
            f'[算法] MoE门控: y=Σ{R("g(x)_i")}·E_i(x), 总参数{R("大")}但每次前向仅激活{R("少数")}专家,稀疏计算控制成本',
            f'[细节] FFN由{R("两")}层全连接网络组成,第一层后接{R("ReLU")}激活,d_ff通常设为{R("4×d_model")}',
            f'[细节] 残差连接将{R("输入")}直接加到子层输出上,缓解深层网络{R("梯度消失")}问题',
            f'[对比] 编码器用{R("双向")}自注意力;解码器用{R("掩码")}自注意力(只关注自身及{R("之前")}位置)',
        ]),
        ("第3章 预训练数据", [
            f'[对比] 质量过滤两类: 基于{R("分类器")}(GPT-3/PaLM用) vs 基于{R("启发式规则")}(BLOOM/Gopher用)',
            f'[细节] 分类器方法用精选文本训练{R("线性分类器")},给类似网页较高分;可能删除{R("方言/口语")}',
            f'[细节] 启发式规则包括: 语言过滤/{R("困惑度")}指标/统计特征({R("标点分布/符号字比/句长")})/关键词',
            f'[细节] GLaM使用{R("特征哈希线性分类器")}(Feature Hash Based Linear Classifier)高效判断',
            f'[细节] LLaMA去重: 拆段落→小写化→{R("SHA-1")}哈希→前{R("64")}位作键判重',
            f'[算法] BPE完整流程: 确定全词词表→每词切为{R("单字符+词尾标记")}→统计相邻对→合并{R("最高频")}→重复',
            f'[对比] WordPiece vs BPE: BPE按{R("频率")}合并;WordPiece按{R("似然增加最多")}的词元对合并',
            f'[数值] 字节级BPE(GPT-2/BART/LLaMA)以{R("字节")}为基本合并符号，改善{R("非ASCII")}分词',
            f'[细节] 未登录词OOV: 传统方法用{R("[UNK]")}默认表示;子词方法将罕见词分解为{R("已知词元")}序列',
            f'[数值] DeepMind训练{R("400+")}个语言模型(70M-16B参数)研究Scaling Law',
            f'[数值] Chinchilla显著优于Gopher({R("280B")})/Jurassic-1({R("178B")})/Megatron-Turing({R("530B")})',
            f'[数值] 千亿级参数模型每次预训练花费{R("数百万元")}，不可能多次迭代训练',
            f'[细节] OPT采用混合{R("RoBERTa+Pile+PushShift Reddit")}数据,因CC过滤过于繁琐',
            f'[案例] 输入"East Stroudsburg Stroudsburg",模型补全了真实{R("姓名/邮箱/电话/地址")}',
            f'[细节] BigScience ROOTS用基于{R("Transformer+机器翻译")}的NER方法处理{R("100+")}种语言',
            f'[数值] GPT-2在Beam Search b={R("32")}时,对特定上下文会进入{R("重复循环")},反复输出相同内容',
            f'[细节] 句子级重复: 重复单词/短语会造成语言模型陷入{R("Repetition Loops")}',
            f'[对比] Unigram与BPE反向: 从{R("大候选集")}开始,迭代删除使似然增加{R("最少")}的词元',
            f'[细节] Unigram用{R("EM算法")}: 每次迭代先找当前最佳切分({R("Viterbi")}),再重新估计词元概率',
            f'[数值] 原始LLaMA词表大小{R("32K")},很多汉字需{R("2-3")}个Byte Token才能拼成完整汉字',
            f'[数值] LLaMA-2训练{R("2万亿")}词元,Baichuan-2训练{R("2.6万亿")}词元,训练文件约{R("10TB")}',
            f'[对比] 开源数据集: Pile({R("825GB")}/22源)/ROOTS({R("1.6T")}token)/RefinedWeb({R("5T")}token)/SlimPajama({R("627B")}token)',
            f'[数值] PaLM记忆率: 训练数据见1次→{R("0.75%")}记忆率; 见500+次→{R(">40%")}记忆率',
            f'[数值] Gopher数据最优配比: {R("10%")}C4+{R("50%")}MassiveWeb+{R("30%")}Books+{R("10%")}News',
            f'[数值] 数据重复影响: 0.1%数据×100次重复→800M模型性能降至{R("400M")}水平',
        ]),
        ("第4章 分布式训练", [
            f'[对比] CPU({R("复杂逻辑")},低延迟少量高性能核心) vs GPU({R("并行计算")},数千小型核心高吞吐)',
            f'[对比] TPU({R("脉动阵列")},硬连线数据流,高能效比) vs NPU({R("MAC乘加阵列")},低功耗实时性)',
            f'[里程碑] TPU由谷歌{R("2016.5")}首次公开,用于AlphaGo; 华为NPU麒麟970 {R("2017")}年发布',
            f'[细节] 集群架构: 服务器→机柜{R("Rack")}→架顶交换机{R("ToR")}→骨干交换机{R("Spine Switch")}',
            f'[细节] 大模型训练通常采用{R("胖树Fat-Tree")}拓扑,试图实现网络带宽{R("无收敛")}',
            f'[细节] 集合通信原语: Broadcast/Scatter/{R("Reduce")}/{R("AllReduce")}/Gather/{R("AllGather")}/AlltoAll',
            f'[对比] 同步训练(等待所有梯度→聚合→更新,慢但{R("稳定")}) vs 异步(收到即更新,快但{R("波动")})',
            f'[数值] H100 HBM带宽{R("3350GB/s")}; PCIe 5.0仅{R("128GB/s")}; 服务器内用{R("NVLink")}互联',
            f'[数值] 训练集群: 8卡A100/H100 SXM终端,{R("400Gb+")}InfiniBand网络',
            f'[数值] LLaMA-2 70B={R("172万")}GPU小时,1024卡A100集群需{R("70")}天',
            f'[数值] ZeRO-1和ZeRO-2通信量{R("无影响")}; ZeRO-3是正常通信量的{R("1.5")}倍',
            f'[数值] 7.5亿参数: FP16推理仅需{R("15GB")},训练需{R("120GB")}(含Adam+激活值)',
            f'[细节] PyTorch内置{R("ZeroRedundancyOptimizer")}封装Adam,step()内存峰值为Adam{R("一半")}',
            f'[数值] GPT-3批次从{R("32K")}逐渐增加到{R("3.2M")}词元',
            f'[细节] DeepSpeed训练: (1)训练前评估{R("困惑度")}→(2)训练循环(前向→损失→梯度→{R("参数更新")})→(3){R("保存模型")}',
            f'[细节] 模型保存: HuggingFace格式({R("from_pretrained")}直接加载)或DeepSpeed {R("Zero Stage 3")}特定格式',
            f'[细节] 激活值检查点: 前向时{R("丢弃")}中间激活,反向时{R("重新计算")},大幅减少内存',
            f'[数值] LR热身{R("0.1%-0.5%")}→峰值{R("5e-5~1e-4")}→余弦衰减至最大值约{R("10%")}直至收敛',
            f'[数值] AdamW超参: β1={R("0.9")}, β2={R("0.95")}, ε={R("1e-8")}; 梯度裁剪阈值{R("1.0")}; 权重衰减率{R("0.1")}',
            f'[对比] 参数服务器: 训练服务器计算梯度→{R("PS聚合更新")}→拉取新参数; 瓶颈在{R("PS带宽")}',
            f'[细节] LLaMA-2全局批次大小{R("4M")}词元; GPT-3批次从{R("32K")}逐渐增加到{R("3.2M")}词元',
        ]),
        ("第5章 有监督微调", [
            f'[算法] 提示学习三阶段: 提示{R("添加")}(拼接模板)→答案{R("搜索")}(找最高分)→答案{R("映射")}(→标签)',
            f'[细节] 软提示Soft Prompt: 不再要求提示是{R("自然语言")},模板有自己{R("可学习")}的参数',
            f'[细节] LIMA数据来源: 高质量{R("网络问答社区")}+Super-Natural Instructions+标注者{R("手写")}',
            f'[细节] LIMA质量过滤: 选靠前回答→统一{R("AI助手")}风格→删过长过短→删{R("第一人称")}→删含链接',
            f'[算法] Self-Instruct四步: 种子{R("175")}条→采样8条(6种子+2模型)引导生成→判断{R("分类/非分类")}→ROUGE-L<0.7',
            f'[对比] 分类任务用{R("输出优先")}(先产生标签再补输入); 非分类用{R("输入优先")}(先产生输入)',
            f'[代码] DeepSpeed-Chat参数: --step({R("1/2/3")}), --deployment-type({R("single_gpu/multi_node")})',
            f'[代码] 数据格式: JSON含{R("prompt")}和{R("chosen")}字段(RM微调还有{R("rejected")}字段)',
            f'[细节] DeepSpeed-Chat默认使用{R("LLaMA-2 7B")},中文可切换为{R("Baichuan 7B")}',
            f'[细节] 训练过程计算{R("困惑度PPL")},PPL随训练进行逐步{R("下降")}',
            f'[细节] DeepSpeed-Chat三大功能: 易用类ChatGPT训练/{R("DeepSpeed-RLHF管道")}(复现InstructGPT)/RLHF系统',
            f'[数值] ROUGE-L过滤阈值{R("0.7")}: 新指令ROUGE-L<0.7才可加入保证{R("多样性")}',
            f'[数值] LoRA: W\'=W0+BA, B(d×r){R("随机")}初始化, A(r×d){R("零")}初始化, scale={R("α/r")}',
            f'[数值] LoRA效果: r=4仅Q/V矩阵, 350GB→{R("35MB")}(万分之一), GPU 1.2TB→{R("350GB")}, 速度提升{R("25%")}',
            f'[对比] AdaLoRA: ΔW=PΓQ, {R("SVD")}参数化, 自适应分配秩(奇异值={R("重要性分数")})',
            f'[对比] QLoRA: {R("NF4")}量化(基于N(0,1)分位数,4bit=16分位), 65B模型仅需{R("48GB")}显存',
            f'[细节] LoRA推理合并: W=W0+BA, {R("无额外")}推理延迟,可直接部署',
            f'[数值] ALiBi斜率序列: m={R("2^(-8/n)")}, {R("2^(-16/n)")},...为{R("几何")}序列,头数越多衰减越快',
            f'[细节] ICL语境学习: 在prompt中提供示例,模型{R("推断")}任务模式; Few-Shot通常K={R("1-64")}个示例',
            f'[数值] 开源指令集: Alpaca({R("52K")})/ShareGPT({R("70K")})/OpenAssistant/{R("FLAN")}',
            f'[细节] NTK-Aware Scaling: {R("非均匀")}缩放不同频率,低频{R("外推")}高频{R("插值")}',
            f'[数值] LIMA仅{R("1000")}条高质量指令数据即可媲美{R("数十倍")}数据的同等规模模型',
        ]),
        ("第6章 强化学习", [
            f'[细节] PPO涉及{R("四")}个模型: 策略/{R("奖励")}/评论(预测累积奖励)/{R("参考")}(SFT备份防极端变化)',
            f'[算法] PPO三步: (1)环境{R("采样")}(策略生成+奖励打分)→(2){R("GAE")}优势估计→(3)参考模型{R("约束")}更新',
            f'[细节] HH-RLHF收集: Anthropic通过{R("Amazon Mechanical Turk")}聊天工具,标注者选一个{R("继续对话")}',
            f'[细节] 有用性数据: 标注者{R("开放式")}对话选更好回答; 无害性: 引诱模型给{R("有害")}回答',
            f'[对比] 有用性和无害性往往{R("对立")}: 过度追求无害→安全但{R("无用")}; 过度有用→可能{R("有害")}',
            f'[算法] 奖励模型: 移除{R("最后非嵌入层")}→叠加{R("线性层")}→输出标量奖励值',
            f'[算法] 总奖励函数: r_total=r(x,y)-{R("η·KL(π_RL‖π_SFT)")},KL散度促进{R("探索")}+防止偏离',
            f'[数值] Summarize数据集: 对比部分{R("17.9万")}条+轴向部分{R("1.5万")}条(Likert量表)',
            f'[对比] Stanford SHP({R("38.5万")})来自{R("Reddit")}自然产生 vs HH-RLHF来自{R("机器生成")}',
            f'[细节] SHP数据: Reddit帖子(问题+两条高赞评论),{R("点赞更多")}者为人类偏爱回复',
            f'[细节] PPO监控推荐综合标准: {R("PPL")}+模型输出{R("长度")}+回复{R("奖励")}',
            f'[对比] RL vs RLHF对比表: 目标({R("环境奖励")}vs{R("人类偏好")})/奖励来源({R("环境反馈")}vs人工标注)/场景',
            f'[对比] RL核心算法{R("DQN/PPO/Q-learning")} vs RLHF核心算法{R("PPO+RM/DPO/RLAIF")}',
            f'[数值] HH-RLHF训练集{R("16.1万")}条; 收集时不记录偏好强度,每对当{R("二选一")}权重相同',
            f'[细节] GPT-4评估时需精心设计{R("提示语")},考虑{R("位置敏感性")}确保公正评价',
            f'[算法] 优势函数: A(s,a)={R("Q(s,a)-V(s)")},衡量动作比{R("平均水平")}好多少; GAE用多步TD折中偏差与方差',
            f'[对比] RL元素→LLM映射: 状态={R("prompt")}, 动作={R("token")}, 策略={R("语言模型")}, 奖励={R("奖励模型")}',
            f'[细节] RLOO({R("leave-one-out")}基线): 降低PPO对评论模型的依赖; GRPO({R("Group Relative")}策略优化): DeepSeek-R1采用',
            f'[数值] WebGPT数据集: {R("1.9万")}条对比数据,用于提升{R("长文档问答")}能力',
            f'[细节] {R("3H")}原则(InstructGPT): {R("Helpfulness")}(有帮助)/{R("Honesty")}(诚实)/{R("Harmlessness")}(无害)',
            f'[细节] Reward Hacking: 模型快速提高奖励但输出{R("无意义/重复")}→增大{R("KL惩罚")}力度→确保奖励缓慢提升',
            f'[数值] 人类偏好标注一致性仅{R("60%-70%")}; 标注过程需控制问题{R("多样性")}和标注标准',
            f'[算法] 模仿学习RM损失: L_rm=L_p+β_rm·{R("L_lm")}, L_lm为{R("自回归")}语言模型损失',
            f'[对比] 三类用户输入: (a){R("Text-Grounded")}文本问答 (b){R("Knowledge-Seeking")}知识查询 (c){R("Creative")}创意生成',
            f'[细节] PPO-Max总结{R("7")}种关键稳定性因素: KL惩罚项/奖励值{R("正则化")}与裁剪/评论模型损失裁剪等',
        ]),
        ("第7章 大语言模型应用", [
            f'[细节] CoT由{R("Google Brain")}提出: 提供解题思路和步骤,模型不仅输出{R("结果")}还输出{R("中间步骤")}',
            f'[算法] Auto-CoT: (1)Sentence-BERT计算问题向量→{R("K-means聚类")}为K簇→(2)每簇选距中心{R("最近")}+Zero-Shot CoT',
            f'[细节] LangChain记忆: {R("读取")}和{R("写入")}两操作; 链在核心逻辑前从记忆读取,执行后{R("保存")}',
            f'[对比] ConversationBufferMemory({R("简单消息列表")}) vs 高级记忆({R("ChatModel")}结合,多对话共享)',
            f'[细节] Agent组件: Agent({R("决策类")})+Tools({R("调用函数")})+Toolkits({R("工具集合3-5个")})+AgentExecutor',
            f'[细节] AgentExecutor还支持{R("Plan-and-execute")}Agent/{R("Baby AGI")}/{R("Auto GPT")}等',
            f'[代码] StdOutCallbackHandler是LangChain内置最基本{R("处理程序")},将所有事件记录到{R("stdout")}',
            f'[算法] 知识库问答5步: 收集→文本{R("提取+分割")}→嵌入+{R("向量库")}→检索+{R("合并输入")}LLM→返回',
            f'[细节] RAG扩展模块: 查询{R("重写/分解")}→检索→{R("重排序")}(Reranker)→压缩→生成; 缓解知识过时/幻觉/{R("本地知识")}缺失',
            f'[细节] 智能体具有{R("创造工具")}能力: 自动编写API调用代码/集成现有工具为更强工具',
            f'[对比] 多智能体优势: (1){R("数量")}优势(分工,各专其能)(2){R("质量")}优势(多观点反馈,减少幻觉)',
            f'[细节] "{R("针锋相对")}"Tit-for-Tat状态下,智能体可从其他智能体获得{R("外部反馈")}纠正扭曲思维',
            f'[细节] 角色扮演能激发LLM内部{R("独特领域知识")},产生比不指定角色时{R("更好")}的答案',
            f'[案例] 行业LLM: BloombergGPT({R("金融")})/ChatLaw({R("法律")})/DISC-MedLLM({R("医疗")})/EduChat({R("教育")})',
            f'[案例] SheetCopilot理解{R("自然语言指令")}自动执行{R("电子表格")}操作',
            f'[细节] 大语言模型推理遵循{R("自回归")}模式: 每次迭代时间确定但{R("迭代次数")}未知,总时间不可预测',
            f'[对比] BERT执行时间{R("确定")}且高度可预测 vs LLM推理{R("迭代次数")}未知导致总时间不可预测',
            f'[细节] FastServe三大目标: (1)低{R("JCT")}(2)高效GPU{R("显存管理")}(3)可扩展{R("分布式")}系统',
            f'[细节] FastServe键值缓存: 主动将{R("低优先级")}作业KV张量转移到{R("主机内存")},动态调整',
            f'[案例] 研究人员利用LLM{R("代码理解")}和{R("调试")}能力发现Linux内核中的未知bug',
            f'[案例] InstructUIE/UniversalNER将数十个{R("信息抽取")}任务用一个LLM实现,超BERT{R("单任务")}效果',
            f'[数值] vLLM由{R("UC Berkeley")}开发,部署于Chatbot Arena和{R("Vicuna")}Demo',
            f'[数值] PagedAttention: 非连续KV缓存,避免{R("60-80%")}内存浪费,吞吐量比HF Transformers高{R("24")}倍',
            f'[细节] KV Cache两阶段: {R("初始化")}(处理prompt,为每层生成KV缓存)→{R("解码")}(仅计算新token的QKV)',
            f'[细节] 由少至多提示(Least-to-Most): 引导模型先{R("分解")}复杂问题为子问题,再{R("逐一")}求解',
            f'[细节] CoT不同人员编写的范例在准确率上存在高达{R("28.2%")}的差异; 改变范例顺序仅产生不到{R("2%")}变化',
            f'[细节] CAMEL框架: {R("双Agent")}(用户+助手)+可选任务明确Agent+{R("评论")}Agent',
            f'[数值] MiniGPT-4架构: EVA-CLIP ViT-G/14+{R("Q-Former")}+Vicuna-{R("13B")}; 预训练{R("20K")}步/4×A100/{R("10h")}',
            f'[细节] 智能体记忆模块: {R("世界知识")}记忆(LLM参数隐式存储)+{R("社会属性")}记忆(外置存储身份和交互历史)',
            f'[数值] Orca: {R("迭代级")}调度,每批仅执行{R("1")}次迭代(1 token/job); {R("FCFS")}策略; 有头部阻塞',
        ]),
        ("第8章 大语言模型评估", [
            f'[细节] AGIEval包含: 中国高考/美国{R("SAT")}/LSAT/{R("数学竞赛")}/司法/{R("公务员")}考试',
            f'[细节] 复杂推理三类: {R("知识推理")}(CSQA/StrategyQA)/{R("符号推理")}(字母连接)/{R("数学推理")}(GSM8K/MATH)',
            f'[数值] StrategyQA: {R("2780")}个评估数据,每个含问题/{R("推理步骤")}/{R("证据段落")}',
            f'[细节] CSQA构造: 从{R("ConceptNet")}过滤边→抽取子图→众包编写→添加{R("干扰概念")}',
            f'[细节] 符号推理评估用{R("ID")}(域内,步骤与训练相同)和{R("OOD")}(域外,步骤更多)测试集',
            f'[数值] LISA: {R("18.3万")}定理+{R("216万")}证明步(Isabelle证明器); miniF2F: {R("488")}道IMO级',
            f'[细节] 具身AI评估环境: {R("VirtualHome")}/ALFRED/BEHAVIOR/{R("Voyager")}/GITM(Minecraft)',
            f'[数值] API-Bank: {R("53")}种API/{R("264")}个对话/{R("568")}个API调用,评估工具使用能力',
            f'[数值] CAIL-Long: 刑事平均{R("916")}汉字,民事{R("1286")}汉字,共{R("112万+110万")}案件',
            f'[数值] LeCaRD: {R("107")}个查询+{R("43000+")}候选案例,来自最高法{R("刑事")}案件',
            f'[细节] Chatbot Arena: 众包匿名对比→用户输入→获{R("两匿名")}回答→投票→提交后{R("显示模型名")}',
            f'[对比] LLMEVAL五种方法: 分项/{R("众包对比")}/{R("公众对比")}/GPT-4分项/GPT-4对比,各有优缺点',
            f'[细节] C-EVAL专业层面参考{R("国家职业资格目录")},选{R("12")}个代表性职业(医生/律师/公务员)',
            f'[代码] pass@k: 生成k个代码样本至少{R("1个通过")}测试的概率,代码评估{R("核心指标")}',
            f'[细节] C-EVAL专业层面参考{R("国家职业资格目录")},选{R("12")}个代表性职业领域(医生/律师/公务员)',
            f'[数值] Chatbot Arena发布{R("33K")}对话数据(2023年4-6月),使用{R("Elo")}评分计算综合分数',
            f'[细节] 具身AI评估: VirtualHome程序步骤step_t=[{R("action_t")}]<object_t,1>(id_t,1)...',
            f'[对比] 人工评估(MOS均分,{R("$80")}/800条) vs LLM-as-a-Judge(GPT-4,快速低成本但有{R("偏差")})',
            f'[数值] 截至2023.8,国内外已发布{R("120+")}个开源/闭源大语言模型',
            f'[数值] HELM框架: {R("42")}场景/{R("59")}指标/{R("16")}核心场景,三维分类(任务/领域/语言)',
            f'[数值] CrowS-Pairs: {R("1508")}例/{R("9")}类偏见(种族/性别/性取向/宗教/年龄/国籍/残疾/外貌/社经)',
            f'[数值] Winogender: {R("120")}人工构建句对,性别名词替换({R("he/she")})检测{R("性别")}偏见',
            f'[细节] LLaMA-2安全风险三类: (1){R("违法犯罪")}(2){R("有害冒犯")}(3){R("不合格建议")}(医疗/金融/法律)',
            f'[细节] 6种指令攻击: 目标{R("劫持")}/提示{R("泄露")}/角色扮演/不安全主题/注入不可见内容/{R("反向")}暴露',
            f'[对比] 红队测试4种方法: (1)零样本(2)随机少样本(3){R("SFT")}监督学习(4){R("A2C")}强化学习最大化有害性',
            f'[数值] MMLU: {R("15858")}题/{R("57")}学科/4选1; 人类Amazon众包{R("34.5%")}vs专业人员约{R("89.8%")}',
            f'[数值] C-EVAL: {R("52")}学科/{R("4")}难度(初中/高中/大学/专业); 大学覆盖{R("13")}个本科专业大类',
            f'[细节] C-EVAL HARD: C-EVAL中非常具有{R("挑战性")}的子集,需要{R("高级推理")}能力',
            f'[算法] BLEU精确率导向: BP·exp(ΣWn·logPn); BP=1(lc>lr)或{R("exp(1-lr/lc)")}惩罚短句',
            f'[算法] ROUGE召回率导向: ROUGE-N({R("n-gram")}匹配)/ROUGE-L({R("LCS最长公共子序列")})',
            f'[数值] 困惑度PPL=2^H=P(s)^(-1/W); 英文n-gram参考范围{R("50-1000")}; 交叉熵{R("6-10")}bit',
            f'[数值] HumanEval({R("164")}题)/MBPP({R("974")}题)/CodeContest(竞赛级)是{R("代码")}评估数据集',
            f'[算法] McNemar检验: χ²=(|B-C|-1)²/(B+C); 显著性水平α={R("0.05")}; B+C<25时用{R("二项式")}精确检验',
            f'[对比] Cohen\'s Kappa({R("2")}人评估) vs Fleiss\' Kappa({R("≥3")}人评估): Pa/Pe多评估者一致性',
            f'[数值] Chatbot Arena胜率: GPT-4 vs GPT-3.5-Turbo胜率{R("79%")}, vs LLaMA-13B胜率{R("94%")}',
        ]),
    ]
    html_parts = []
    for ch_title, items in chapters:
        html_parts.append(f'<div class="fbch"><b class="fbcht">{ch_title} · 多类型速记</b>')
        for item in items:
            html_parts.append(f'<div class="fbi">{item}</div>')
        html_parts.append('</div>')
    return '\n'.join(html_parts)

# ===== Formula section with LaTeX (v7: +11 new formulas) =====
def build_formula_section():
    """All formulas in LaTeX, rendered by KaTeX. v7 adds 11 missing formulas."""
    lines = [
        # --- Core Transformer ---
        ('位置编码',
         r'\text{PE}(pos,2i)=\sin\!\Big(\frac{pos}{10000^{2i/d}}\Big),\;\text{PE}(pos,2i\!+\!1)=\cos\!\Big(\frac{pos}{10000^{2i/d}}\Big)'),
        ('自注意力',
         r'\text{Attn}(Q,K,V)=\text{softmax}\!\Big(\frac{QK^{T}}{\sqrt{d_k}}\Big)V'),
        ('多头注意力',
         r'h_i=\text{Attn}(XW_i^Q,XW_i^K,XW_i^V);\;\text{Multi}=\text{Concat}\cdot W^O'),
        ('FFN',
         r'\text{FFN}(x)=\text{ReLU}(xW_1+b_1)W_2+b_2'),
        ('残差连接',
         r'x_{l+1}=x_l+f(x_l)'),
        ('RMSNorm',
         r'\frac{a}{\text{RMS}(a)}\gamma,\;\text{RMS}(a)=\sqrt{\tfrac{1}{d}\sum a_i^2}'),
        ('SwiGLU',
         r'\text{FFN}(x)=(xW_1\odot\sigma(xW_3))W_2'),
        ('Swish',
         r'x\cdot\sigma(\beta x);\;\beta\!\to\!0\text{ 线性},\;\beta\!\to\!\infty\text{ ReLU}'),
        ('RoPE',
         r'f(q,m)=q\cdot e^{im\theta},\;\theta=10000^{-2i/d}'),
        # --- NEW: RoPE位置插值 ---
        ('RoPE插值',
         r"f(x,m\!\cdot\!\tfrac{L_{\text{pre}}}{L_{\text{ext}}})\;\text{扩展至}32768"),
        # --- GPT / Training ---
        ('GPT训练',
         r'\mathcal{L}=\sum_t\log P(w_t|w_{&lt;t});\;\text{微调:}\;\mathcal{L}=\mathcal{L}_{sft}+\lambda\mathcal{L}_{pre}'),
        ('混合精度',
         r'\text{参数}2\Phi+\text{梯度}2\Phi+\text{Adam}(4\!+\!4\!+\!4)\Phi=\mathbf{16\Phi}'),
        ('ZeRO',
         r'S_1\!\to\!\tfrac{1}{4},\;S_2\!\to\!\tfrac{1}{8},\;S_3\!\to\!\text{线性}'),
        # --- NEW: 梯度累积 ---
        ('梯度累积',
         r'\text{eff\_batch}=\text{micro\_batch}\times\text{accum\_steps}'),
        # --- NEW: 分布式Softmax ---
        ('分布式Softmax',
         r'\text{sm}(x)_p=\frac{e^{x_p-\max}}{\sum_q e^{x_q-\max}}\;\text{(三次通信)}'),
        # --- LoRA family ---
        ('LoRA',
         r"W'=W_0+BA,\;\text{scale}=\tfrac{\alpha}{r};\;350\text{GB}\!\to\!35\text{MB}"),
        ('AdaLoRA/QLoRA',
         r'\Delta W=P\Gamma Q;\;\text{QLoRA NF4:}\;\mathcal{N}(0,1)\;\text{4bit量化}'),
        # --- ALiBi + NEW slope ---
        ('ALiBi',
         r'\text{softmax}(qK^T - m|i-j|);\;m\!=\!2^{-8/n},2^{-16/n},\ldots'),
        # --- RL ---
        ('奖励模型',
         r'\mathcal{L}=-\log\sigma(r_w-r_l);\;\mathcal{L}_{rm}=\mathcal{L}_p+\beta\mathcal{L}_{lm}'),
        ('策略梯度',
         r'\nabla J=\mathbb{E}[R\cdot\nabla\!\log p];\;+\text{bl}:(R-b)\nabla\!\log p'),
        ('折扣回报',
         r"\sum_{t'}\gamma^{t'-t}r_{t'}"),
        ('GAE',
         r'A=Q-V,\;\mathcal{L}_V=\mathbb{E}[(V_\phi-Q)^2]'),
        # --- NEW: 重要性采样 ---
        ('重要性采样',
         r'\mathbb{E}_{p}[f]=\mathbb{E}_{q}\!\big[f\cdot\tfrac{p}{q}\big];\;r=\tfrac{p_\theta}{p_{\theta\'}}'),
        ('PPO-Clip',
         r'\min\!\big(r\cdot A,\;\text{clip}(r,1\!-\!\varepsilon,1\!+\!\varepsilon)\cdot A\big)'),
        ('PPO-Penalty',
         r"\mathbb{E}[r\cdot A]-\beta\cdot\text{KL}(\theta\|\theta')"),
        # --- NEW: 红队A2C ---
        ('红队A2C',
         r'\alpha\cdot\mathcal{L}_{A2C}+(1\!-\!\alpha)\text{KL}(p_r\|p_{\text{init}})'),
        # --- Evaluation ---
        ('困惑度PPL',
         r'2^{H}=P(s)^{-1/W}'),
        ('BLEU',
         r'\text{BP}\cdot\exp\!\Big(\sum_n W_n\log P_n\Big)'),
        # --- NEW: BLEU BP ---
        ('BLEU-BP',
         r'\text{BP}=1(l_c\!&gt;\!l_r);\;\text{BP}=e^{1-l_r/l_c}(l_c\!\leq\! l_r)'),
        ('ROUGE-N',
         r'\frac{\sum\text{匹配n-gram}}{\sum Y\text{中n-gram}}'),
        # --- NEW: ROUGE-L ---
        ('ROUGE-L',
         r'R=\tfrac{\text{LCS}}{|Y|},P=\tfrac{\text{LCS}}{|\hat{Y}|},F=\tfrac{(1+\beta^2)RP}{R+\beta^2 P}'),
        ('P/R/F1',
         r'P=\frac{TP}{TP+FP},\;R=\frac{TP}{TP+FN},\;F_1=\frac{2PR}{P+R}'),
        # --- NEW: MSLE + MedAE ---
        ('MSLE/MedAE',
         r'\tfrac{1}{n}\sum(\log(y_i\!+\!1)-\log(\hat{y}_i\!+\!1))^2;\;\text{med}(|y_i\!-\!\hat{y}_i|)'),
        ('McNemar',
         r'\chi^2=\frac{(|B-C|-1)^2}{B+C};\;B\!+\!C\!&lt;\!25\text{:二项式精确}'),
        # --- NEW: Fleiss' Kappa Pa/Pe ---
        ('Fleiss\'Kappa',
         r'P_a=\tfrac{\sum n_{ij}(n_{ij}-1)}{|X|n(n-1)},\;P_e=\sum\!\big(\tfrac{n_j}{|X|n}\big)^2'),
        ('Cohen\'s Kappa',
         r'\kappa=\frac{P_a-P_c}{1-P_c}'),
        # --- Tokenization ---
        ('BPE',
         r'\text{全词}\to\text{单字符+&lt;/w&gt;}\to\text{统计频率}\to\text{合并最高频}\to\text{重复}'),
        ('WordPiece',
         r'\text{Score}=\frac{\text{Cnt(pair)}}{\text{Cnt}(a)\cdot\text{Cnt}(b)}'),
        ('PagedAttn',
         r'\text{非连续KV},\;-60\text{–}80\%\text{浪费},\;\text{吞吐}\uparrow 24\times'),
        # --- NEW from PPT gap analysis ---
        ('TD误差',
         r'\delta_t=r_t+\gamma V_\phi(s_{t+1})-V_\phi(s_t)'),
        ('GAE指数形式',
         r'A_t^{\text{GAE}}=\sum_{k=0}^{\infty}(\gamma\lambda)^k\delta_{t+k}'),
        ('k步优势',
         r'A_t^{(k)}=\sum_{l=0}^{k-1}\gamma^l\delta_{t+l};\;\text{MC:高方差无偏,TD:低方差有偏}'),
        ('RL总奖励',
         r'r_{\text{total}}=r(x,y)-\eta\cdot\text{KL}(\pi_{RL}\|\pi_{SFT})'),
        ('LSH(Reformer)',
         r'h(x)=\text{argmax}([xR;-xR]);\;R\in\mathbb{R}^{d_k\times b/2}'),
        # --- NEW from review outline ---
        ('链式法则',
         r'P(w_1,\ldots,w_m)=\prod_{i=1}^{m}P(w_i|w_1,\ldots,w_{i-1})'),
        ('注意力复杂度',
         r'\text{Time/Space:}\;O(n^2 d);\;\text{长序列瓶颈}'),
        ('奖励排序',
         r'R(x,y_w)>R(x,y_l);\;\text{偏好:chosen vs rejected}'),
        ('MoE',
         r'y=\sum_i g(x)_i\cdot E_i(x);\;\text{总参大,激活少}'),
        ('分布式速度',
         r'v_{\text{train}}\propto v_{\text{dev}}\times N_{\text{dev}}\times\eta_{\text{acc}}'),
    ]
    html_parts = []
    for label, latex in lines:
        html_parts.append(f'<b class="hd">{label}:</b> \\({latex}\\)')
    return '\n'.join(html_parts)

# ===== Numerical section (v7: +12 new entries) =====
def build_numerical_section():
    """Numerical references. v7 adds all missing values from PPT."""
    lines = [
        ('模型规模', r'\text{BERT}{\sim}110\text{M},\;\text{GPT-1}{\sim}117\text{M},\;\text{T5}{\sim}11\text{B},\;\text{GPT-2}{\sim}1.5\text{B},\;\text{GPT-3}{\sim}175\text{B},\;\text{PaLM}{\sim}540\text{B}'),
        ('GPT-4 MoE', r'1.8\text{T总参},\;16\text{专家}\!\times\!1110\text{亿},\;\text{路由}2,\;\text{共享}550\text{亿}'),
        ('Chinchilla', r'70\text{B}+1.4\text{T词元(等比缩放,优于Gopher 280B)}'),
        ('GPT-3资源', r'\text{FP32}=700\text{GB},\;\text{FP16}=350\text{GB},\;314\text{ZFLOPs}'),
        ('H100', r'80\text{GB HBM},\;2000\text{TFLOPs FP16},\;\text{HBM}\;3350\text{GB/s},\;\text{SRAM}\;228\text{KB/块}'),
        ('网络带宽', r'\text{PCIe5.0:}128\text{GB/s},\;\text{IB:}200\text{-}400\text{Gbps},\;\text{DGX:}1.6\text{Tb},\;\text{HGX:}3.2\text{Tb}'),
        ('通信量', r'128\text{副本}\times 700\text{GB}=89.6\text{TB/迭代};\;\text{1-bit Adam:}1/5\text{通信}'),
        ('LLaMA训练', r'7\text{B}=82\text{K},\;13\text{B}=135\text{K},\;65\text{B}=1022\text{K GPUh};\;70\text{B}=172\text{万GPUh}'),
        ('LLaMA-2', r'2\text{T词元训练};\;\text{Baichuan-2:}2.6\text{T词元}'),
        ('数据流水线', r'\text{CC:}45\text{TB}\!\to\!570\text{GB}\!\to\!500\text{B词元};\;\text{Wiki:}3.4\text{epoch},\;\text{CC:}0.44\text{epoch}'),
        ('训练阶段', r'\text{预训练:}1000+\text{GPU/月};\;\text{SFT/RM/RL:}1\text{-}100\text{GPU/天}'),
        ('Embedding TP', r'64000\!\times\!5120\!\times\!4\text{B}\approx\!1.25\text{GB}\!\times\!2(\text{含梯度})\!=\!2.5\text{GB}'),
        ('Adam', r'\beta_1=0.9,\;\beta_2=0.95,\;\varepsilon=10^{-8},\;\text{clip}=1.0,\;\text{wd}=0.1'),
        ('LR调度', r'\text{热身}0.1\text{–}0.5\%,\;\text{峰值}5\text{e-}5,\;\text{余弦衰减至}10\%'),
        ('批次大小', r'\text{LLaMA-2:}4\text{M词元};\;\text{GPT-3:}32\text{K}\!\to\!3.2\text{M词元}'),
        ('数据重复', r'0.1\%\!\times\!100\text{次}\!\to\!\text{性能减半};\;\text{HP复制11次}\!\to\!3\%\text{重复严重影响}'),
        ('MQA', r'\text{共享K/V,仅保留Q,约}5\%\text{数据微调即可}'),
        ('FlashAttn', r'\text{S/P/O不入全局内存};\;\text{SRAM 228KB/块}\gg\text{HBM速度}'),
        ('PPL参考', r'\text{英文n-gram:}50\text{–}1000;\;\text{交叉熵:}6\text{–}10\text{ bit}'),
        # --- NEW from PPT gap analysis ---
        ('DGX GH200', r'256\text{ GH200芯片},\;144\text{TB共享内存},\;2023.5'),
        ('模型增速', r'\text{AlexNet→PaLM(9年):}56\times\text{/18个月}'),
        ('RLHF数据', r'\text{HH-RLHF:}161\text{K},\;\text{SHP:}385\text{K},\;\text{WebGPT:}19\text{K},\;\text{Summ:}179\text{K}'),
        ('MMLU', r'15858\text{题},\;57\text{学科},\;\text{众包}34.5\%\text{ vs 专家}89.8\%'),
        ('AGIEval', r'\text{GPT-4: SAT数学}95\%,\;\text{高考英语}92.5\%,\;\text{超人平均}'),
        ('Chatbot Arena', r'\text{GPT-4 vs GPT-3.5:}79\%\text{胜率};\;\text{vs LLaMA-13B:}94\%'),
        ('数学数据', r'\text{GSM8K:}8500,\;\text{MATH:}12500,\;\text{LISA:}183\text{K定理},\;\text{miniF2F:}488'),
        ('MiniGPT-4', r'\text{预训练:}20\text{K步/4×A100/10h};\;\text{微调:}400\text{步/1GPU/7min}'),
        ('LIMA', r'1000\text{条高质量指令≈数十倍数据效果(Less is More)}'),
        ('标注一致率', r'\text{人类偏好标注一致性:}60\text{–}70\%'),
        ('PaLM记忆率', r'\text{见1次:}0.75\%;\;\text{见500+次:}&gt;40\%'),
        ('Gopher数据混合', r'\text{最优:}10\%\text{C4}+50\%\text{MassiveWeb}+30\%\text{Books}+10\%\text{News}'),
        ('安全评估', r'6000+\text{场景},\;2800\text{攻击样本},\;6\text{攻击类型};\;\text{CrowS-Pairs:}1508\text{例/9类偏见}'),
        ('<30B无需TP', r'\text{模型}<30\text{B参数无需张量并行,现有框架即可}'),
        ('Self-Instruct', r'175\text{条种子指令→自动扩展};\;\text{LLaMA默认}2048\text{token窗口}'),
    ]
    html_parts = []
    for label, latex in lines:
        html_parts.append(f'<b class="hd">{label}:</b> \\({latex}\\)')
    return '\n'.join(html_parts)

# ===== Calculation problems with step-by-step solutions =====
def build_calc_section():
    """12 key calculation problems from PPT with detailed solutions."""
    calcs = [
        ("【1】混合精度训练内存(GPT-3 175B)",
         "Adam+FP16混合精度: 参数2Φ+梯度2Φ+Adam动量(4+4+4)Φ=16Φ字节(Φ=参数量)。"
         "GPT-3: 175B×16B=2800GB。7.5B模型: FP16推理仅15GB, 训练需120GB(含Adam状态)。"),
        ("【2】ZeRO优化器内存缩减",
         "原始16Φ。S1(Adam分区): 16Φ/N→占1/4; S2(+梯度分区): 8Φ/N→占1/8; "
         "S3(+参数分区): 与数据并行度线性。例: 8卡S3→每卡仅需16Φ/8=2Φ。"),
        ("【3】LoRA参数量与显存节省",
         "原矩阵W(d×d=4096²=16.7M参数)。LoRA: B(d×r=4096×4)+A(r×d=4×4096)=32768参数,"
         "仅为原来的0.2%。检查点: 350GB→35MB(万分之一)。GPU: 1.2TB→350GB,速度↑25%。"),
        ("【4】GPT-3分布式训练通信量",
         "175B参数FP32=175G×4B=700GB/副本。128副本×700GB=89.6TB/迭代。"
         "IB带宽200-400Gbps=25-50GB/s, 传输89.6TB需~1800-3600秒(30-60分钟)。"),
        ("【5】Embedding层张量并行内存",
         "词表64000×维度5120×4B(FP32)=1,310,720,000B≈1250MB。"
         "含反向梯度: 1250MB×2=2.5GB。张量并行切到8卡: 每卡仅~313MB。"),
        ("【6】ROUGE-1/ROUGE-2计算",
         "候选Y_hat='a dog is in the garden', 参考Y='there is a dog in the garden'。"
         "ROUGE-1: 匹配词{a,dog,is,in,the,garden}=6, Y中1-gram=7, R=6/7≈0.857。"
         "ROUGE-2: 匹配bigram{is a,a dog,dog is,in the,the garden}=5, Y中2-gram=6, R=5/6≈0.833。"),
        ("【7】BLEU惩罚因子BP",
         "机器译文长度lc, 最短参考译文长度lr。若lc>lr: BP=1(无惩罚)。"
         "若lc≤lr: BP=exp(1-lr/lc)。例: lc=15,lr=20 → BP=exp(1-20/15)=exp(-0.333)≈0.717。"),
        ("【8】McNemar检验(模型对比)",
         "模型1正确/模型2错误=B, 反向=C。χ²=(|B-C|-1)²/(B+C)。"
         "例: B=25,C=15 → χ²=(10-1)²/40=81/40=2.025。B+C<25时用二项式精确检验。"),
        ("【9】困惑度PPL计算",
         "PPL=2^H=P(s)^(-1/W)。例: 句子概率P(s)=0.001, 词数W=10。"
         "H=-log₂(0.001)/10=9.97/10≈1.0。PPL=2^1.0=2.0。英文参考: PPL 50-1000。"),
        ("【10】数据并行梯度计算",
         "M设备, 批次N样本, 每设备N/M个。设备i计算本地梯度Gi。"
         "平均梯度: G_avg=(1/M)ΣGi。AllReduce通信量: 2(M-1)/M × 模型参数量。"),
        ("【11】梯度累积等效batch",
         "显存受限时的micro_batch=4, accumulation_steps=8。"
         "eff_batch=4×8=32。等效于batch_size=32的全量训练, 但显存仅需batch_size=4的量。"),
        ("【12】CommonCrawl数据流水线",
         "原始45TB→质量过滤后570GB(保留1.27%)→子词切分后~500B词元。"
         "训练300B词元: Wiki平均3.4 epoch, CC仅0.44 epoch, Books2仅0.43 epoch。"),
    ]
    html_parts = []
    for title, solution in calcs:
        html_parts.append(f'<b class="hd">{title}</b><br>{solution}<br>')
    return '\n'.join(html_parts)

# ===== Essay/论述题 section with万能结构 + 对比题 =====
def build_essay_section():
    """论述题万能结构 + 10道高频对比题."""
    # 论述万能结构
    struct = [
        ("技术流程类", "先写目标→再写输入数据→再写关键模块→再写训练/推理过程→最后写优势和风险。"),
        ("对比类", "先给定义→再从对象、输入、过程、优点、缺点、适用场景六个角度对比。"),
        ("系统设计类", "先画流程: 输入→预处理→核心模型/检索/工具→输出→评估反馈。再说明每模块作用。"),
        ("风险分析类", "从数据风险、模型风险、推理风险、安全风险、评估风险五个角度展开。"),
    ]
    # 10道对比题
    comps = [
        ("空间智能 vs 具身智能", "数字空间/API调用/文本生成 vs 物理载体/多模态传感器/实时闭环; 部署2-4周 vs 研发6-12+月"),
        ("Encoder vs Decoder", "双向自注意力(理解任务,BERT) vs 掩码自注意力+交叉注意力(生成任务,GPT); 前者用全上下文,后者仅用左上下文"),
        ("预训练 vs SFT vs RLHF", "语言分布建模(无标注,海量数据) vs 指令遵循(少量标注) vs 偏好对齐(人类反馈); 计算量: 极大→小→小"),
        ("数据并行 vs 模型并行", "切数据(每设备完整模型,分样本) vs 切模型(每设备部分模型,分参数); DP受显存限制,MP可训超大模型"),
        ("流水线并行 vs 张量并行", "按层切分(设备间串行,有气泡) vs 按矩阵切分(设备间并行,需通信); PP适合层间,TP适合层内"),
        ("全量微调 vs LoRA", "更新全部参数(显存大,效果最好) vs 低秩增量ΔW=BA(r<<d,仅训练0.2%参数); LoRA: 350GB→35MB"),
        ("RAG vs 微调", "知识注入上下文(检索→拼接→生成,实时更新) vs 知识注入参数(训练固化,更新需重训); RAG无需重训但受上下文窗口限制"),
        ("Chain vs Agent", "固定流程(预定义步骤链,确定性高) vs 动态决策(LLM自主选择工具和路径,灵活性高); Chain适合标准流程,Agent适合开放任务"),
        ("BLEU vs ROUGE", "BLEU=精确率导向(机器翻译,n-gram精确率+BP惩罚) vs ROUGE=召回率导向(文本摘要,n-gram召回率); BLEU关注译文质量,ROUGE关注信息覆盖"),
        ("人工评估 vs LLM-as-Judge", "人工: 准确但贵慢($80/800条) vs LLM评估: 快速低成本但有偏差; Chatbot Arena用Elo排名结合人类偏好"),
    ]
    html_parts = []
    html_parts.append('<b class="hd">论述万能结构:</b><br>')
    for title, desc in struct:
        html_parts.append(f'<b class="r">{title}:</b> {desc}<br>')
    html_parts.append('<br><b class="hd">10道高频对比题:</b><br>')
    for title, desc in comps:
        html_parts.append(f'<b class="r">{title}:</b> {desc}<br>')
    return '\n'.join(html_parts)

# ===== Doc2 table LaTeX helper =====
def latex_cell(text):
    if not text or len(text) < 3:
        return text
    math_indicators = ['=', 'Σ', 'σ', '√', 'log', 'softmax', 'exp', 'sigmoid',
                       'θ', 'γ', 'β', 'α', 'ε', 'λ', 'φ', 'Φ', 'Δ', '∇',
                       'softmax', 'PPL', 'BLEU', 'ROUGE', 'F1', 'KL']
    has_math = any(ind in text for ind in math_indicators)
    if re.search(r'[A-Z]\s*[=/(]', text):
        has_math = True
    if has_math:
        t = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return f'\\({t}\\)'
    return text

def build_d2_table_latex(table_data):
    if not table_data:
        return ''
    rows_html = []
    for ri, row in enumerate(table_data):
        cells = ''
        for c in row:
            if ri == 0:
                cells += f'<td>{c}</td>'
            else:
                cells += f'<td>{latex_cell(c)}</td>'
        rows_html.append(f'<tr>{cells}</tr>')
    return '<table>' + ''.join(rows_html) + '</table>'

# ===== Main generator =====
def gen_html(data, d2, d2_tables):
    mcq_ans = data['mcq_ans_full']
    fill_qs = data['fill_qs']
    fill_ans = data['fill_ans_full']
    defn_ans = data['defn_ans']
    short_ans = data['short_ans']
    app_ans = data['app_ans']
    checklist = data['checklist']
    
    # MCQ answer grid
    mcq_ans_dict = {n: l for n, l, _ in mcq_ans}
    max_n = max(mcq_ans_dict.keys()) if mcq_ans_dict else 40
    mcq_grid = '\u2002'.join(str(i) + mcq_ans_dict.get(i, '?') for i in range(1, max_n + 1))
    
    # MCQ: answer + key theory
    mcq_html = ''
    for n, letter, body in mcq_ans:
        key_theory = ''
        explanation = ''
        for line in body.split('  '):
            line = line.strip()
            if line.startswith('关键理论') or line.startswith('关键'):
                key_theory = line
            elif line.startswith('解释'):
                explanation = line
            elif not key_theory:
                key_theory = line
            else:
                explanation += ' ' + line
        mcq_html += f'<b>{n}.</b><span class="r">{letter}</span> {key_theory} {explanation}<br>'
    
    # Fill-in-blank
    fill_html = ''
    for n, q in fill_qs:
        ans_tuple = next(((a, b) for nn, a, b in fill_ans if nn == n), ('', ''))
        a_text, a_body = ans_tuple
        qr = re.sub(r'_{2,}', f'<span class="r">{a_text}</span>', q, count=1)
        if qr == q:
            fill_html += f'<b>{n}.</b>{q} → <span class="r">{a_text}</span>'
        else:
            fill_html += f'<b>{n}.</b>{qr}'
        if a_body:
            fill_html += f' <i>{a_body[:120]}</i>'
        fill_html += '<br>'
    
    # Additional PPT-based fill-in-blank questions (26-45)
    extra_fill = [
        (26, 'GPT-3训练时Wikipedia语料平均训练____次epoch,CommonCrawl仅____次。', '3.4 / 0.44'),
        (27, 'OPT使用____块NVIDIA A100 80G GPU，训练近____个月。', '992 / 2'),
        (28, 'BLOOM训练花费____个月，使用____个计算节点，每节点8块A100，总计____GPU。', '3.5 / 48 / 384'),
        (29, 'LLaMA与标准Transformer有____处关键不同：RMSNorm+____+RoPE+GQA。', '四 / SwiGLU'),
        (30, 'FlashAttention将Q/K/V分块加载到____中计算，中间结果S/P/O____写回全局内存。', 'SRAM / 不'),
        (31, 'BPE完整流程：确定全词词表→每词切为____→统计相邻对→合并____→重复。', '单字符+词尾标记 / 最高频'),
        (32, 'Chinchilla定律：模型大小和训练词元应____缩放，模型加倍→数据____。', '等比例 / 加倍'),
        (33, 'LLaMA-2 70B训练____万GPU小时，用1024卡A100集群需____天。', '172 / 70'),
        (34, 'ZeRO-1和ZeRO-2对通信量____影响；ZeRO-3通信量是正常的____倍。', '无 / 1.5'),
        (35, 'LoRA: W\'=W0+BA，B矩阵____初始化，A矩阵____初始化，推理时____额外延迟。', '随机 / 零 / 无'),
        (36, 'PPO涉及____个模型：策略模型/____模型/评论模型/____模型。', '四 / 奖励 / 参考'),
        (37, 'RLHF中总奖励函数：r_total=r(x,y)-____·KL散度，KL散度促进____并防止偏离。', 'η / 探索'),
        (38, 'HH-RLHF训练集数据总量____条，通过____聊天工具收集。', '16.1万 / Amazon Mechanical Turk'),
        (39, 'CoT由____提出：提供解题思路和步骤，模型不仅输出结果还输出____。', 'Google Brain / 中间步骤'),
        (40, 'LangChain智能体组件：Agent(____)+Tools(____)+Toolkits(工具集合)+AgentExecutor。', '决策类 / 调用函数'),
        (41, 'AGIEval包含：中国高考/美国____/LSAT/____竞赛/司法/公务员考试。', 'SAT / 数学'),
        (42, 'MMLU数据集：____题/____学科/4选1；人类Amazon众包仅34.5%。', '15858 / 57'),
        (43, 'C-EVAL：____学科/____难度(初中/高中/大学/专业)；专业层面参考国家职业资格目录。', '52 / 4'),
        (44, 'Chatbot Arena：众包匿名对比评估，使用____评分计算综合分数；发布____对话数据。', 'Elo / 33K'),
        (45, 'pass@k：生成k个代码样本至少____通过测试的概率，是代码评估的____指标。', '1个 / 核心'),
    ]
    start_n = 26
    for n, q, ans in extra_fill:
        qr = q.replace('____', f'<span class="r">{ans}</span>', 1)
        fill_html += f'<b>{n}.</b>{qr}<br>'
    
    defn_html = '<br>'.join(f'<b>{i+1}.</b>{a}' for i, a in enumerate(defn_ans))
    
    short_html = ''
    for n, title, body in short_ans:
        short_html += f'<b>{n}.</b>{title}<br><span class="r">答:</span> {body}<br>'
    
    app_html = ''
    for n, title, body in app_ans:
        app_html += f'<b>{n}.</b>{title}<br><span class="r">答:</span> {body}<br>'
    
    cl_html = '<br>'.join(t for t in checklist)
    formulas = build_formula_section()
    numerics = build_numerical_section()
    knowledge_map = build_knowledge_map()
    fillblank_html = build_fillblank_review()
    calc_html = build_calc_section()
    essay_html = build_essay_section()
    
    # Doc2 tables
    tbl_cmp = ''
    if d2_tables and len(d2_tables) > 0:
        tbl_cmp = build_d2_table_latex(d2_tables[0])
    tbl_frm = ''
    if d2_tables and len(d2_tables) > 1:
        tbl_frm = build_d2_table_latex(d2_tables[1])
    
    d2_cmp = ''
    for k in d2:
        if '对比' in k:
            for typ, t in d2[k]: d2_cmp += t + '<br>'
    d2_frm = ''
    for k in d2:
        if '公式' in k:
            for typ, t in d2[k]: d2_frm += t + '<br>'
    d2_struct = ''
    for k in d2:
        if '论述' in k or '万能' in k:
            for typ, t in d2[k]:
                if typ == 'h3': d2_struct += f'<b class="hd">{t}</b> '
                else: d2_struct += t + ' '
            break
    d2_mem = ''
    for k in d2:
        if '背诵' in k or '临考' in k:
            for typ, t in d2[k]: d2_mem += t + '<br>'
            break

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>具身大模型理论 · 速查表（完整版·双页A4·LaTeX公式）</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"
  onload="renderMathInElement(document.body, {{
    delimiters: [
      {{left: '\\\\(', right: '\\\\)', display: false}}
    ],
    throwOnError: false,
    fontSize: '0.80em',
    trust: true
  }});"></script>
<style>
@page {{ size: A4 portrait; margin: 2mm 2mm 2mm 2mm; }}
* {{ margin:0; padding:0; box-sizing:border-box; }}
html,body {{
  width:210mm;
  font-family:"Microsoft YaHei","SimSun",sans-serif;
  font-size:4.2pt; line-height:1.22; color:#111;
}}
.page {{
  width:206mm; height:294mm;
  page-break-after: always;
  overflow:hidden;
}}
.page:last-child {{ page-break-after: auto; }}
.cols6 {{
  column-count:5; column-gap:2mm;
  column-rule:0.3pt solid #ddd;
  height:292mm;
}}
h1 {{
  font-size:5.8pt; text-align:center;
  column-span:all; margin-bottom:0.5mm;
  padding:0.4mm 0; background:#16213e; color:#fff;
  letter-spacing:0.8pt;
}}
h2 {{
  font-size:4.5pt; background:#16213e; color:#fff;
  padding:0.25mm 0.6mm; margin:0.6mm 0 0.3mm;
  break-inside:avoid; break-after:avoid;
}}
.hd {{ color:#16213e; font-weight:bold; }}
.r {{ color:#c0392b; font-weight:bold; }}
sup,sub {{ font-size:3.2pt; }}
i {{ font-style:normal; color:#555; }}
table {{
  width:100%; border-collapse:collapse;
  font-size:3.5pt; margin:0.2mm 0;
  break-inside:avoid;
}}
td {{ border:0.2pt solid #aaa; padding:0.1mm 0.2mm; }}
tr:first-child td {{ background:#16213e; color:#fff; font-weight:bold; }}
.sec {{ break-inside:avoid; margin-bottom:0.15mm; }}
.f {{ font-size:3.9pt; line-height:1.18; }}
/* KaTeX overrides */
.katex {{ font-size: 0.80em !important; }}
.katex .mord {{ white-space: nowrap; }}
/* Page 2 mind map styles */
.cols5 {{
  column-count:5; column-gap:2mm;
  column-rule:0.3pt solid #ddd;
}}
.cols5-page2 {{
  column-count:5; column-gap:2mm;
  column-rule:0.3pt solid #ddd;
  height:291mm;
}}
.cols5-upper {{
  height:100mm;
  overflow:hidden;
}}
.cols5-lower {{
  height:190mm;
  border-top:0.5pt dashed #999;
  padding-top:0.5mm;
  overflow:hidden;
}}
.fbch {{
  margin-bottom:0.8mm;
  border-left:1.5pt solid #c0392b;
  padding-left:0.8mm;
}}
.fbcht {{
  display:block;
  font-size:4.4pt;
  color:#fff;
  background:#c0392b;
  padding:0.2mm 0.6mm;
  margin:-0.1mm 0 0.3mm -0.8mm;
}}
.fbi {{
  font-size:3.9pt;
  line-height:1.2;
  margin-bottom:0.15mm;
  padding-left:0.3mm;
}}
.ch {{
  margin-bottom:1.2mm;
  border-left:1.8pt solid #16213e;
  padding-left:1mm;
}}
.cht {{
  display:block;
  font-size:5.2pt;
  color:#fff;
  background:#16213e;
  padding:0.35mm 1mm;
  margin:-0.2mm 0 0.5mm -1mm;
}}
.sc {{
  font-size:4.4pt;
  line-height:1.26;
  margin-bottom:0.4mm;
  padding-left:0.5mm;
  border-left:0.6pt solid #c0392b;
}}
.shd {{
  color:#16213e;
  font-size:4.4pt;
}}
.map-legend {{
  column-span:all;
  font-size:3.2pt;
  text-align:center;
  margin-bottom:0.5mm;
  color:#666;
}}
</style>
</head>
<body>

<!-- ========== PAGE 1: 答案 + 公式 + 题目 ========== -->
<div class="page">
<div class="cols6">
<h1>具身大模型理论 · 速查表 第1页（公式·数值·答案）</h1>

<h2>A. 公式大全（43个核心公式）</h2>
<div class="sec f">{formulas}</div>

<h2>B. 关键数值（34项）</h2>
<div class="sec f">{numerics}</div>

<h2>C. 对比 &amp; 公式表</h2>
<div class="sec">{tbl_cmp}{tbl_frm}</div>
<div class="sec f">{d2_cmp}{d2_frm}</div>

<h2>D. 论述万能结构 &amp; 对比题</h2>
<div class="sec f">{essay_html}</div>

<h2>E. 临考背诵</h2>
<div class="sec f">{d2_mem}</div>

<h2>F. 选择题答案速查</h2>
<div class="sec f">{mcq_grid}</div>

<h2>G. 选择题答案+解释</h2>
<div class="sec f">{mcq_html}</div>

<h2>H. 填空题（45题·含答案+解释）</h2>
<div class="sec f">{fill_html}</div>

<h2>I. 名词解释</h2>
<div class="sec f">{defn_html}</div>

<h2>J. 简答题（完整解答）</h2>
<div class="sec f">{short_html}</div>

<h2>K. 综合应用题（完整解答）</h2>
<div class="sec f">{app_html}</div>

<h2>L. 计算题（12道·含详细步骤）</h2>
<div class="sec f">{calc_html}</div>

<h2>M. 考前速记</h2>
<div class="sec f" style="color:#c0392b">{cl_html}</div>

</div>
</div>

<!-- ========== PAGE 2: 知识图谱 + 填空速记 ========== -->
<div class="page">
<h1 style="font-size:5pt; text-align:center; padding:0.3mm 0; background:#16213e; color:#fff; letter-spacing:0.6pt; margin-bottom:0.3mm;">具身大模型理论 · 速查表 第2页（知识图谱 + 章节填空速记）</h1>
<div class="cols5-page2">
{knowledge_map}
<hr style="column-span:all; border:none; border-top:0.5pt dashed #999; margin:1mm 0;">
{fillblank_html}
</div>
</div>

</body>
</html>'''
    return html

# ===== MAIN =====
d1_paras, d1_tables = extract_all(r'C:\Users\Administrator\Downloads\测试习题与参考答案完整版.docx')
d2_paras, d2_tables = extract_all(r'C:\Users\Administrator\Downloads\考点速查.docx')

h1s = find_headings(d1_paras, 'Heading 1')
h2s = find_headings(d1_paras, 'Heading 2')
all_h = sorted(h1s + h2s, key=lambda x: x[0])

h1_map = {}
for idx, s, t in h1s:
    if '单项选择' in t: h1_map['mcq_q'] = idx
    elif '填空' in t: h1_map['fill_q'] = idx
    elif '名词解释' in t and '参考' not in t: h1_map['defn_q'] = idx
    elif '简答' in t and '参考' not in t: h1_map['short_q'] = idx
    elif '综合应用' in t and '参考' not in t: h1_map['app_q'] = idx
    elif '参考' in t or '答案' in t: h1_map['ans'] = idx
    elif '速背' in t or '速记' in t or '清单' in t: h1_map['checklist'] = idx

h2_map = {}
for idx, s, t in h2s:
    if '单项选择' in t: h2_map['mcq_a'] = idx
    elif t == '' or t.strip() == '': h2_map['fill_a'] = idx
    elif '名词解释' in t: h2_map['defn_a'] = idx
    elif '简答' in t: h2_map['short_a'] = idx
    elif '综合' in t: h2_map['app_a'] = idx

def end_of(start_idx):
    for h_idx, _, _ in all_h:
        if h_idx > start_idx: return h_idx
    return len(d1_paras)

data = {
    'mcq_ans_full': parse_mcq_answers_full(get_section(d1_paras, h2_map['mcq_a'], end_of(h2_map['mcq_a']))),
    'fill_qs': parse_fill_questions(get_section(d1_paras, h1_map['fill_q'], end_of(h1_map['fill_q']))),
    'fill_ans_full': parse_fill_answers_full(get_section(d1_paras, h2_map['fill_a'], end_of(h2_map['fill_a']))),
    'defn_ans': parse_defn_answers(get_section(d1_paras, h2_map['defn_a'], end_of(h2_map['defn_a']))),
    'short_ans': parse_essay_full(get_section(d1_paras, h2_map['short_a'], end_of(h2_map['short_a']))),
    'app_ans': parse_essay_full(get_section(d1_paras, h2_map['app_a'], end_of(h2_map['app_a']))),
    'checklist': [t for s, t in d1_paras[h1_map['checklist']+1:] if t],
}

d2_items = [(s, t) for s, t in d2_paras if t]
d2_sections = {}; cur = None
for s, t in d2_items:
    if 'Heading 1' in s: continue
    if 'Heading 2' in s: cur = t; d2_sections[cur] = []
    elif 'Heading 3' in s:
        if cur: d2_sections[cur].append(('h3', t))
    else:
        if cur and t: d2_sections[cur].append(('p', t))

html = gen_html(data, d2_sections, d2_tables)
out = r'C:\Users\Administrator\.qoderwork\workspace\mq0nryuisw8500xo\outputs\具身大模型理论_速查表_完整版.html'
os.makedirs(os.path.dirname(out), exist_ok=True)
with open(out, 'w', encoding='utf-8') as f:
    f.write(html)
print(f'Output: {out}')
print(f'Size: {os.path.getsize(out)} bytes')
print(f'Formulas: 38 (was 27, +11 new)')
print(f'Numerical entries: 19 (was 11, +8 new)')
print(f'Knowledge map: updated with all PPT values')
print(f'MCQ answers: {len(data["mcq_ans_full"])}')
print(f'Fill answers: {len(data["fill_ans_full"])}')
print(f'Short answers: {len(data["short_ans"])}')
print(f'App answers: {len(data["app_ans"])}')
