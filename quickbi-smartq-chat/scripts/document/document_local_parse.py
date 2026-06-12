#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用文档文本提取工具（纯本地执行，不依赖任何外部 API）。

支持 PDF、DOC、DOCX、XLSX、XLS、CSV 及常用图片格式的文本提取。
支持单文件、多文件、文件夹递归扫描，并行处理。

技术栈：
- PDF: PyMuPDF 直接提取文本，扫描件使用 Tesseract OCR（逐页渲染为图片后识别）
- 图片: Tesseract OCR 本地识别（中英文混合）
- Word: python-docx 提取（.docx），libreoffice 转换 + OCR（.doc）
- Excel: openpyxl/xlrd 直接提取
- CSV: pandas 读取（自动检测编码）

系统依赖：
- Tesseract OCR: 
  * macOS: brew install tesseract tesseract-lang
  * Linux: sudo apt install tesseract-ocr tesseract-ocr-chi-sim
  * Windows: 从 https://github.com/UB-Mannheim/tesseract/wiki 下载安装
- LibreOffice (可选，用于 .doc 格式转换):
  * macOS: brew install --cask libreoffice
  * Linux: sudo apt install libreoffice
  * Windows: 从 https://www.libreoffice.org 下载安装
- Python 包: pytesseract, Pillow, PyMuPDF, python-docx, openpyxl, xlrd, pandas

用法：
    from document_local_parse import extract_text
    
    # 提取单个文件
    text = extract_text("/path/to/file.pdf")
    
    # 批量提取
    texts = extract_text(["/path/to/file1.pdf", "/path/to/file2.docx"])
    
    # 扫描文件夹（递归）
    texts = extract_text("/path/to/folder/")
    
    # 命令行使用
    python document_local_parse.py file1.pdf file2.png --json
    python document_local_parse.py /path/to/folder/ --output result.json
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import platform
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# ---------------------------------------------------------------------------
# 格式分类
# ---------------------------------------------------------------------------

PDF_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.webp'}
WORD_EXTENSIONS = {'.doc', '.docx'}
EXCEL_EXTENSIONS = {'.xls', '.xlsx'}
CSV_EXTENSIONS = {'.csv'}

ALL_SUPPORTED_EXTENSIONS = (
    PDF_EXTENSIONS | IMAGE_EXTENSIONS | WORD_EXTENSIONS | EXCEL_EXTENSIONS | CSV_EXTENSIONS
)

# 输出目录
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"

# 最大并行数
MAX_WORKERS = 10


# ---------------------------------------------------------------------------
# 跨平台工具函数
# ---------------------------------------------------------------------------

def get_libreoffice_path() -> str:
    """
    获取跨平台的 libreoffice 可执行文件路径。
    
    Returns:
        libreoffice 可执行文件路径
    """
    system = platform.system()
    
    if system == 'Windows':
        # Windows 常见安装路径
        win_paths = [
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]
        for p in win_paths:
            if p.exists():
                return str(p)
        return 'libreoffice'  # fallback: 假设在 PATH 中
    elif system == 'Darwin':  # macOS
        mac_path = Path('/Applications/LibreOffice.app/Contents/MacOS/soffice')
        if mac_path.exists():
            return str(mac_path)
        return 'libreoffice'  # fallback: 假设通过 brew 安装且在 PATH 中
    else:  # Linux
        return 'libreoffice'


def check_tesseract_available() -> bool:
    """
    检查 Tesseract OCR 是否可用。
    
    Returns:
        True 如果可用，否则 False
    """
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        system = platform.system()
        if system == 'Windows':
            msg = "请从 https://github.com/UB-Mannheim/tesseract/wiki 下载并安装，添加到 PATH"
        elif system == 'Darwin':
            msg = "运行: brew install tesseract tesseract-lang"
        else:  # Linux
            msg = "运行: sudo apt install tesseract-ocr tesseract-ocr-chi-sim (Debian/Ubuntu) 或 sudo yum install tesseract (CentOS)"
        print(f"[OCR检查] ⚠ Tesseract 未安装或不可用: {msg}", flush=True)
        return False


def safe_remove_temp_file(file_path: str):
    """
    安全删除临时文件（处理 Windows 文件占用问题）。
    
    Args:
        file_path: 临时文件路径
    """
    try:
        if os.path.exists(file_path):
            os.unlink(file_path)
    except PermissionError:
        # Windows 文件被占用时延迟删除
        if platform.system() == 'Windows':
            import atexit
            print(f"[临时文件] 警告: 文件被占用，将在程序退出时删除: {file_path}", flush=True)
            atexit.register(lambda: os.path.exists(file_path) and os.unlink(file_path))
        else:
            raise


# ---------------------------------------------------------------------------
# 文件扫描工具
# ---------------------------------------------------------------------------

def scan_files(path: Union[str, Path], recursive: bool = True) -> List[Path]:
    """
    扫描文件路径，支持单文件、多文件、文件夹递归。
    
    Args:
        path: 文件路径或文件夹路径
        recursive: 是否递归子文件夹
        
    Returns:
        支持的文件路径列表
    """
    path = Path(path)
    
    if path.is_file():
        # 单个文件
        ext = path.suffix.lower()
        if ext in ALL_SUPPORTED_EXTENSIONS:
            return [path]
        else:
            print(f"[扫描] 跳过不支持的文件: {path.name} ({ext})", flush=True)
            return []
    
    elif path.is_dir():
        # 文件夹扫描
        files = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in ALL_SUPPORTED_EXTENSIONS:
                    files.append(file_path)
        
        print(f"[扫描] 在 {path} 中找到 {len(files)} 个支持的文件", flush=True)
        return files
    
    else:
        print(f"[扫描] 路径不存在: {path}", flush=True)
        return []


# ---------------------------------------------------------------------------
# PDF 文本提取
# ---------------------------------------------------------------------------

def extract_pdf_text(file_path: Path, use_ocr_fallback: bool = True) -> str:
    """
    提取 PDF 文本。优先使用 PyMuPDF，如果是扫描件则降级到 OCR。
    
    Args:
        file_path: PDF 文件路径
        use_ocr_fallback: 是否在文本提取失败时降级到 OCR
        
    Returns:
        提取的文本内容
    """
    import fitz  # PyMuPDF
    
    text = ""
    try:
        doc = fitz.open(str(file_path))
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            text += page_text
        
        doc.close()
        
        # 检查是否成功提取到足够的文本
        # 如果文本太少（< 50 字符），可能是扫描件，需要 OCR
        if len(text.strip()) < 50 and use_ocr_fallback:
            print(f"[PDF提取] 警告: 本地提取文本较少 ({len(text.strip())} 字符)，可能是扫描件，尝试 OCR...", flush=True)
            return _ocr_fallback(file_path)
        
        print(f"[PDF提取] 成功提取 {len(text.strip())} 字符", flush=True)
        return text.strip()
        
    except Exception as e:
        print(f"[PDF提取] 本地提取失败: {e}", flush=True)
        if use_ocr_fallback:
            print("[PDF提取] 降级到 OCR 识别...", flush=True)
            return _ocr_fallback(file_path)
        raise


# ---------------------------------------------------------------------------
# 图片 OCR 提取
# ---------------------------------------------------------------------------

def extract_image_text(file_path: Path, use_ocr: bool = True) -> str:
    """
    提取图片中的文本（OCR）。
    
    Args:
        file_path: 图片文件路径
        use_ocr: 是否使用 OCR（目前仅支持 QuickBI OCR API）
        
    Returns:
        识别的文本内容
    """
    if not use_ocr:
        print(f"[图片提取] 警告: 图片文件必须使用 OCR 提取，但 use_ocr=False", flush=True)
        return ""
    
    print(f"[图片提取] 使用 OCR 识别: {file_path.name}", flush=True)
    return _ocr_fallback(file_path)


# ---------------------------------------------------------------------------
# Word 文档提取
# ---------------------------------------------------------------------------

def extract_word_text(file_path: Path) -> str:
    """
    提取 Word 文档文本。
    
    Args:
        file_path: Word 文件路径（.doc 或 .docx）
        
    Returns:
        提取的文本内容
    """
    ext = file_path.suffix.lower()
    
    if ext == '.docx':
        return _extract_docx(file_path)
    elif ext == '.doc':
        # .doc 格式较老，尝试使用 antiword 或转换为 .docx
        return _extract_doc(file_path)
    else:
        raise ValueError(f"不支持的 Word 格式: {ext}")


def _extract_docx(file_path: Path) -> str:
    """提取 .docx 文件文本。"""
    from docx import Document
    
    try:
        doc = Document(str(file_path))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        text = "\n".join(paragraphs)
        print(f"[Word提取] 成功提取 {len(text.strip())} 字符", flush=True)
        return text.strip()
        
    except Exception as e:
        print(f"[Word提取] 提取失败: {e}", flush=True)
        raise


def _extract_doc(file_path: Path) -> str:
    """
    提取 .doc 文件文本。
    策略：尝试使用 libreoffice 转换为 .docx
    """
    # 方法1: 尝试使用 libreoffice 转换为 .docx
    try:
        import subprocess
        
        libreoffice_cmd = get_libreoffice_path()
        print(f"[Word提取] 使用 libreoffice: {libreoffice_cmd}", flush=True)
        
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdocx = Path(tmpdir) / f"{file_path.stem}.docx"
            
            # Windows 下如果路径包含空格，需要使用 shell=True
            use_shell = platform.system() == 'Windows' and ' ' in libreoffice_cmd
            
            # 使用 libreoffice 转换
            result = subprocess.run(
                [
                    libreoffice_cmd, '--headless', '--convert-to', 'docx',
                    '--outdir', str(tmpdir), str(file_path)
                ],
                capture_output=True,
                text=True,
                timeout=30,
                shell=use_shell
            )
            
            if result.returncode == 0 and tmpdocx.exists():
                print(f"[Word提取] 已将 .doc 转换为 .docx", flush=True)
                return _extract_docx(tmpdocx)
            else:
                print(f"[Word提取] libreoffice 转换失败: {result.stderr}", flush=True)
            
    except FileNotFoundError:
        print("[Word提取] 警告: 未找到 libreoffice，尝试其他方法...", flush=True)
    except Exception as e:
        print(f"[Word提取] libreoffice 转换失败: {e}", flush=True)
    
    # 方法2: 尝试使用 textract
    try:
        import textract
        text = textract.process(str(file_path)).decode('utf-8')
        print(f"[Word提取] 使用 textract 提取 {len(text.strip())} 字符", flush=True)
        return text.strip()
    except FileNotFoundError:
        print("[Word提取] 警告: 未安装 textract，.doc 文件提取受限", flush=True)
    except Exception as e:
        print(f"[Word提取] textract 提取失败: {e}", flush=True)
    
    # 方法3: 降级到 OCR
    print("[Word提取] 降级到 OCR 识别...", flush=True)
    return _ocr_fallback(file_path)


# ---------------------------------------------------------------------------
# Excel 文档提取
# ---------------------------------------------------------------------------

def extract_excel_text(file_path: Path) -> str:
    """
    提取 Excel 文档文本。
    
    Args:
        file_path: Excel 文件路径（.xls 或 .xlsx）
        
    Returns:
        提取的文本内容（表格格式）
    """
    ext = file_path.suffix.lower()
    
    if ext == '.xlsx':
        return _extract_xlsx(file_path)
    elif ext == '.xls':
        return _extract_xls(file_path)
    else:
        raise ValueError(f"不支持的 Excel 格式: {ext}")


def _extract_xlsx(file_path: Path) -> str:
    """提取 .xlsx 文件文本。"""
    import openpyxl
    
    try:
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        sheets_text = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            
            for row in ws.iter_rows(values_only=True):
                # 过滤空行
                if any(cell is not None and str(cell).strip() for cell in row):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    rows_text.append(row_text)
            
            if rows_text:
                sheets_text.append(f"### 工作表: {sheet_name}\n" + "\n".join(rows_text))
        
        wb.close()
        text = "\n\n".join(sheets_text)
        print(f"[Excel提取] 成功提取 {len(text.strip())} 字符", flush=True)
        return text.strip()
        
    except Exception as e:
        print(f"[Excel提取] 提取失败: {e}", flush=True)
        raise


def _extract_xls(file_path: Path) -> str:
    """提取 .xls 文件文本。"""
    import xlrd
    
    try:
        wb = xlrd.open_workbook(str(file_path))
        sheets_text = []
        
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            rows_text = []
            
            for row_idx in range(ws.nrows):
                row = ws.row_values(row_idx)
                if any(str(cell).strip() for cell in row):
                    row_text = " | ".join(str(cell) for cell in row)
                    rows_text.append(row_text)
            
            if rows_text:
                sheets_text.append(f"### 工作表: {sheet_name}\n" + "\n".join(rows_text))
        
        text = "\n\n".join(sheets_text)
        print(f"[Excel提取] 成功提取 {len(text.strip())} 字符", flush=True)
        return text.strip()
        
    except Exception as e:
        print(f"[Excel提取] 提取失败: {e}", flush=True)
        raise


# ---------------------------------------------------------------------------
# CSV 文件提取
# ---------------------------------------------------------------------------

def extract_csv_text(file_path: Path) -> str:
    """
    提取 CSV 文件文本。
    
    Args:
        file_path: CSV 文件路径
        
    Returns:
        提取的文本内容
    """
    import pandas as pd
    
    try:
        # 尝试多种编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin1']:
            try:
                df = pd.read_csv(str(file_path), encoding=encoding)
                text = df.to_string(index=False)
                print(f"[CSV提取] 成功提取 {len(text.strip())} 字符 (编码: {encoding})", flush=True)
                return text.strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("无法解码 CSV 文件，尝试了 utf-8, gbk, gb2312, latin1")
        
    except Exception as e:
        print(f"[CSV提取] 提取失败: {e}", flush=True)
        raise


# ---------------------------------------------------------------------------
# OCR 降级方案
# ---------------------------------------------------------------------------

def _ocr_fallback(file_path: Path) -> str:
    """
    使用 Tesseract OCR 进行本地识别（纯本地，不依赖任何外部 API）。
    
    Args:
        file_path: 文件路径
        
    Returns:
        OCR 识别的文本内容
    """
    import tempfile
    
    # 检查 Tesseract 是否可用
    if not check_tesseract_available():
        raise RuntimeError(
            "Tesseract OCR 未安装，无法进行 OCR 识别。"
            "请安装后重试，或设置 use_ocr_fallback=False 禁用 OCR 降级。"
        )
    
    try:
        print(f"[OCR降级] 使用 Tesseract OCR 本地识别: {file_path.name}", flush=True)
        
        # 对于 PDF 文件，使用 PyMuPDF 将每页转为图片后再 OCR
        if file_path.suffix.lower() in PDF_EXTENSIONS:
            return _ocr_pdf_with_tesseract(file_path)
        
        # 对于图片文件，直接使用 Tesseract OCR
        elif file_path.suffix.lower() in IMAGE_EXTENSIONS:
            return _ocr_image_with_tesseract(file_path)
        
        # 对于 .doc 等其他文件，先转 PDF 再 OCR
        else:
            return _ocr_other_file_with_tesseract(file_path)
            
    except Exception as e:
        print(f"[OCR降级] OCR 识别失败: {e}", flush=True)
        raise RuntimeError(f"OCR 降级失败: {e}")


def _ocr_pdf_with_tesseract(file_path: Path) -> str:
    """
    使用 Tesseract OCR 识别 PDF 文件（逐页渲染为图片后 OCR）。
    """
    import fitz  # PyMuPDF
    from PIL import Image
    import pytesseract
    import io
    
    print(f"[OCR-PDF] 开始逐页 OCR: {file_path.name}", flush=True)
    
    doc = fitz.open(str(file_path))
    text_parts = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # 将页面渲染为高分辨率图片（300 DPI）
        zoom = 300 / 72  # 300 DPI
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        
        # 转换为 PIL Image
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        
        # 使用 Tesseract OCR（中英文混合）
        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        
        if text.strip():
            text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text.strip()}")
            print(f"[OCR-PDF] 第 {page_num + 1} 页识别成功: {len(text.strip())} 字符", flush=True)
    
    doc.close()
    
    full_text = "\n\n".join(text_parts)
    print(f"[OCR-PDF] 总共识别 {len(text_parts)} 页，{len(full_text)} 字符", flush=True)
    
    return full_text


def _ocr_image_with_tesseract(file_path: Path) -> str:
    """
    使用 Tesseract OCR 识别图片文件。
    """
    from PIL import Image
    import pytesseract
    
    print(f"[OCR-图片] 开始识别: {file_path.name}", flush=True)
    
    # 打开图片
    img = Image.open(str(file_path))
    
    # 对于 .jpg 文件，统一转为 .jpeg 扩展名（Tesseract 支持更好）
    if file_path.suffix.lower() == '.jpg':
        import tempfile
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.jpeg', delete=False) as tmp:
                tmp_path = tmp.name
            img.save(tmp_path, 'JPEG')
            img = Image.open(tmp_path)
        finally:
            # 安全删除临时文件
            if tmp_path:
                safe_remove_temp_file(tmp_path)
    
    # 使用 Tesseract OCR（中英文混合）
    text = pytesseract.image_to_string(img, lang='chi_sim+eng')
    
    if text.strip():
        print(f"[OCR-图片] 识别成功: {len(text.strip())} 字符", flush=True)
        return text.strip()
    else:
        raise RuntimeError("OCR 未能识别到任何文本")


def _ocr_other_file_with_tesseract(file_path: Path) -> str:
    """
    对其他文件格式（如 .doc）进行 OCR。
    策略：先尝试转换为 PDF，然后使用 PDF OCR。
    """
    import fitz
    import subprocess
    import tempfile
    
    print(f"[OCR-其他] 尝试转换并识别: {file_path.name}", flush=True)
    
    # 尝试使用 libreoffice 转换为 PDF
    with tempfile.TemporaryDirectory() as tmpdir:
        tmppdf = Path(tmpdir) / f"{file_path.stem}.pdf"
        
        try:
            libreoffice_cmd = get_libreoffice_path()
            print(f"[OCR-其他] 使用 libreoffice: {libreoffice_cmd}", flush=True)
            
            # Windows 下如果路径包含空格，需要使用 shell=True
            use_shell = platform.system() == 'Windows' and ' ' in libreoffice_cmd
            
            result = subprocess.run(
                [
                    libreoffice_cmd, '--headless', '--convert-to', 'pdf',
                    '--outdir', str(tmpdir), str(file_path)
                ],
                capture_output=True,
                text=True,
                timeout=60,
                shell=use_shell
            )
            
            if result.returncode == 0 and tmppdf.exists():
                print(f"[OCR-其他] 转换成功，开始 OCR", flush=True)
                return _ocr_pdf_with_tesseract(tmppdf)
            else:
                print(f"[OCR-其他] 转换失败: {result.stderr}", flush=True)
        except FileNotFoundError:
            print("[OCR-其他] 未找到 libreoffice", flush=True)
        except Exception as e:
            print(f"[OCR-其他] 转换失败: {e}", flush=True)
        
        # 如果转换失败，尝试直接用 PyMuPDF 打开
        try:
            doc = fitz.open(str(file_path))
            doc.close()
            print(f"[OCR-其他] PyMuPDF 可以直接打开，使用 PDF OCR", flush=True)
            return _ocr_pdf_with_tesseract(file_path)
        except Exception as e:
            print(f"[OCR-其他] PyMuPDF 也无法打开: {e}", flush=True)
            raise RuntimeError(f"无法处理文件类型: {file_path.suffix}")


# ---------------------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------------------

def extract_text(
    file_path: Union[str, Path, List[Union[str, Path]]],
    use_ocr_fallback: bool = True,
    output_dir: Optional[Union[str, Path]] = None,
    save_json: bool = False
) -> Union[str, Dict[str, str], List[Dict[str, str]]]:
    """
    统一文本提取接口。
    
    Args:
        file_path: 文件路径、文件夹路径或文件路径列表
        use_ocr_fallback: 是否在本地提取失败时降级到 OCR
        output_dir: JSON 输出目录（默认 OUTPUT_DIR）
        save_json: 是否保存 JSON 结果文件
        
    Returns:
        - 单个文件: 返回提取的文本字符串
        - 文件列表: 返回 [{"file": "fileName", "parsedText": "text"}] 列表
        
    Examples:
        >>> # 单个文件
        >>> text = extract_text("invoice.pdf")
        
        >>> # 批量提取
        >>> texts = extract_text(["file1.pdf", "file2.docx", "image.png"])
        
        >>> # 扫描文件夹
        >>> results = extract_text("/path/to/folder/", save_json=True)
    """
    # 处理单个文件/文件夹路径
    if isinstance(file_path, (str, Path)):
        path = Path(file_path)
        
        # 如果是文件夹，扫描所有文件
        if path.is_dir():
            files = scan_files(path, recursive=True)
            return _extract_files_parallel(files, use_ocr_fallback, output_dir, save_json)
        
        # 单个文件
        return _extract_single_file(path, use_ocr_fallback)
    
    # 处理文件列表
    return _extract_files_parallel(file_path, use_ocr_fallback, output_dir, save_json)


def _extract_single_file(file_path: Path, use_ocr_fallback: bool = True) -> str:
    """
    提取单个文件的文本。
    
    Args:
        file_path: 文件路径
        use_ocr_fallback: 是否降级到 OCR
        
    Returns:
        提取的文本
    """
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    if not file_path.is_file():
        raise ValueError(f"路径不是文件: {file_path}")
    
    ext = file_path.suffix.lower()
    
    if ext not in ALL_SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件格式: {ext}\n"
            f"支持的格式: {', '.join(sorted(ALL_SUPPORTED_EXTENSIONS))}"
        )
    
    print(f"\n{'='*60}", flush=True)
    print(f"提取文件: {file_path.name}", flush=True)
    print(f"{'='*60}", flush=True)
    
    # 根据文件类型选择提取方法
    if ext in PDF_EXTENSIONS:
        return extract_pdf_text(file_path, use_ocr_fallback)
    elif ext in IMAGE_EXTENSIONS:
        return extract_image_text(file_path, use_ocr_fallback)
    elif ext in WORD_EXTENSIONS:
        return extract_word_text(file_path)
    elif ext in EXCEL_EXTENSIONS:
        return extract_excel_text(file_path)
    elif ext in CSV_EXTENSIONS:
        return extract_csv_text(file_path)
    else:
        raise ValueError(f"未实现的文件类型: {ext}")


def _extract_files_parallel(
    files: List[Union[str, Path]],
    use_ocr_fallback: bool = True,
    output_dir: Optional[Union[str, Path]] = None,
    save_json: bool = False
) -> List[Dict[str, str]]:
    """
    并行提取多个文件的文本。
    
    Args:
        files: 文件路径列表
        use_ocr_fallback: 是否降级到 OCR
        output_dir: JSON 输出目录
        save_json: 是否保存 JSON 文件
        
    Returns:
        [{"file": "fileName", "parsedText": "text"}] 列表
    """
    results = []
    total = len(files)
    
    print(f"\n{'='*60}", flush=True)
    print(f"[并行提取] 开始处理 {total} 个文件 (最大并行数: {MAX_WORKERS})", flush=True)
    print(f"{'='*60}", flush=True)
    
    def process_file(file_path: Union[str, Path]) -> Dict[str, str]:
        """处理单个文件的包装函数"""
        fp = Path(file_path)
        try:
            text = _extract_single_file(fp, use_ocr_fallback)
            return {
                "file": fp.name,
                "parsedText": text
            }
        except Exception as e:
            print(f"[错误] 提取失败 {fp.name}: {e}", flush=True)
            return {
                "file": fp.name,
                "parsedText": f"[ERROR] {e}"
            }
    
    # 使用线程池并行处理
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        # 提交所有任务
        future_to_file = {executor.submit(process_file, fp): fp for fp in files}
        
        # 收集结果
        completed = 0
        for future in as_completed(future_to_file):
            completed += 1
            result = future.result()
            results.append(result)
            print(f"\n[进度] {completed}/{total} 完成: {result['file']}", flush=True)
    
    # 按文件名排序
    results.sort(key=lambda x: x['file'])
    
    print(f"\n{'='*60}", flush=True)
    print(f"[并行提取] 全部完成，成功处理 {len(results)} 个文件", flush=True)
    print(f"{'='*60}", flush=True)
    
    # 保存 JSON 结果
    if save_json or output_dir:
        save_results_json(results, output_dir)
    
    return results


def save_results_json(
    results: List[Dict[str, str]],
    output_dir: Optional[Union[str, Path]] = None
) -> Path:
    """
    保存提取结果为 JSON 文件。
    
    Args:
        results: 提取结果列表
        output_dir: 输出目录（默认 OUTPUT_DIR）
        
    Returns:
        JSON 文件路径
    """
    import time
    
    if output_dir is None:
        output_dir = OUTPUT_DIR
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 生成文件名（带时间戳）
    timestamp = int(time.time())
    json_path = output_dir / f"extract_results_{timestamp}.json"
    
    # 写入 JSON
    json_path.write_text(
        json.dumps(results, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )
    
    print(f"\n[保存] JSON 结果已保存到: {json_path}", flush=True)
    print(f"[保存] 共 {len(results)} 个文件的结果", flush=True)
    
    return json_path


# ---------------------------------------------------------------------------
# 命令行入口
# ---------------------------------------------------------------------------

def main():
    """命令行入口：提取文件文本并输出。"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="通用文档文本提取工具，支持 PDF/Word/Excel/CSV/图片",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 提取单个文件
  python document_local_parse.py invoice.pdf

  # 批量提取多个文件
  python document_local_parse.py file1.pdf file2.docx image.png

  # 扫描文件夹（递归）
  python document_local_parse.py /path/to/folder/

  # 禁用 OCR 降级（仅本地提取）
  python document_local_parse.py invoice.pdf --no-ocr

  # 输出 JSON 结果到默认 output 目录
  python document_local_parse.py /path/to/folder/ --json

  # 输出 JSON 结果到指定目录
  python document_local_parse.py /path/to/folder/ --json --output-dir /custom/output/
        """
    )
    parser.add_argument(
        "files",
        nargs='+',
        help="要提取的文件路径或文件夹路径（可多个）"
    )
    parser.add_argument(
        "--no-ocr",
        action="store_true",
        help="禁用 OCR 降级，仅使用本地提取"
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default=None,
        help="JSON 输出目录（默认为 output/）"
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="保存 JSON 格式结果文件"
    )
    parser.add_argument(
        "--workspace-dir",
        default=None,
        help="用户工作目录路径"
    )

    args = parser.parse_args()

    # 设置工作目录（必须在任何 read_config() 之前）
    if args.workspace_dir:
        from common.config_loader import set_workspace_dir
        set_workspace_dir(args.workspace_dir)
    
    use_ocr = not args.no_ocr
    
    # 收集所有要处理的文件
    all_files = []
    for path_str in args.files:
        path = Path(path_str)
        if path.is_dir():
            # 文件夹：递归扫描
            files = scan_files(path, recursive=True)
            all_files.extend(files)
        elif path.is_file():
            # 单文件
            ext = path.suffix.lower()
            if ext in ALL_SUPPORTED_EXTENSIONS:
                all_files.append(path)
            else:
                print(f"[扫描] 跳过不支持的文件: {path.name} ({ext})", flush=True)
        else:
            print(f"[扫描] 路径不存在: {path}", flush=True)
    
    if not all_files:
        print("[错误] 没有找到可处理的文件", flush=True)
        sys.exit(1)
    
    print(f"\n[准备] 共找到 {len(all_files)} 个文件待处理", flush=True)
    
    # 执行并行提取
    results = extract_text(
        all_files,
        use_ocr_fallback=use_ocr,
        output_dir=args.output_dir,
        save_json=args.json
    )
    
    # 输出结果摘要
    print(f"\n{'='*60}", flush=True)
    print("[结果摘要]", flush=True)
    print(f"{'='*60}", flush=True)
    
    for result in results:
        status = "成功" if not result['parsedText'].startswith('[ERROR]') else "失败"
        text_len = len(result['parsedText'])
        print(f"  {result['file']}: {status} ({text_len} 字符)", flush=True)
    
    print(f"\n总计: {len(results)} 个文件", flush=True)


if __name__ == "__main__":
    main()
